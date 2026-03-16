from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import locale

from models.docTextHelper import resolve_doc_type_key, get_cfg, get_text, get_style, apply_style, apply_alignment, styled_run

def createIndividualWordDocMultiYear(individual, output_file_path, doc_text_config=None):
    if individual[0]['language'] == 'EN':
        createIndividualWordDocEN(individual, output_file_path, doc_text_config)
    if individual[0]['language'] == 'FR':
        createIndividualWordDocFR(individual, output_file_path, doc_text_config)


locale.setlocale(locale.LC_ALL, 'fr_CA.UTF-8')

def format_currency(amount):
    # Convert the amount to a float and format it as a string
    return f"{locale.format_string('%.2f', amount, grouping=True).replace('.', ',')} $"

# Function to set default font and size for the entire document
def set_default_font(doc, font_name, font_size):
    styles = doc.styles['Normal']
    font = styles.font
    font.name = font_name
    font.size = Pt(font_size)
    # Remove default paragraph spacing so sections don't have large gaps
    styles.paragraph_format.space_before = Pt(0)
    styles.paragraph_format.space_after = Pt(0)

# Function that adds a hyperlink to a paragraph.
def add_hyperlink(paragraph, text, url):
    # Create the w:hyperlink tag and add required attributes
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    italic = OxmlElement('w:i')
    rPr.append(italic)
    # Apply properties to the run
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

# Function to create a new document and set global styles
def createIndividualWordDocEN(individual, output_file_path, doc_text_config=None):
    doc = Document()
    set_default_font(doc, "Calibri", 10)

    cfg = get_cfg(doc_text_config, resolve_doc_type_key('EN', is_couple=False, is_multiyear=True))

    ind_title = individual[0]['title']

    isMailQC = any(person.get('isMailQC', False) for person in individual)
    isNewcomer = any(person.get('isNewcomer', False) for person in individual)

    summaries = {person['year']: person['summary'] for person in individual}

    years = sorted(summaries.keys(), key=int)


    # Add section_1 with formatted years
    section_1(doc, summaries[years[0]], ind_title, years, isMailQC, cfg=cfg)

    for year in years:

        return_summary = summaries[year]

        para = doc.add_paragraph()
        run = para.add_run(f"{year}")
        run.underline = True

        tax_return(doc, return_summary, year, years[-1], cfg=cfg)

        if "gst_amounts" in return_summary and return_summary["gst_amounts"]:
            gst_credit(doc, return_summary, isNewcomer, year, years[-1], years[0], cfg=cfg)

        if "ecgeb_amounts" in return_summary and return_summary["ecgeb_amounts"]:
            ecgeb_credit(doc, return_summary, isNewcomer, year, years[-1], years[0], cfg=cfg)

        if "solidarity_amounts" in return_summary and return_summary["solidarity_amounts"]:
            solidarity_credit(doc, return_summary, year, years[-1], cfg=cfg)

        if "ccb_amounts" in return_summary and return_summary["ccb_amounts"]:
            child_benefit(doc, return_summary, year, years[-1], cfg=cfg)

        if "family_allowance_amounts" in return_summary and return_summary["family_allowance_amounts"]:
            quebec_family_allowance(doc, return_summary, year, years[-1], cfg=cfg)

        # Only add carryforward amounts for the last year
        if year == years[-1] and "carryforward_amounts" in return_summary and return_summary["carryforward_amounts"]:
            carryforward_amounts(doc, return_summary, cfg=cfg)

    conclusion(doc, isMailQC, cfg=cfg)
    doc.save(output_file_path)


# Section 1: Title and Header
def section_1(doc, primary_summary, ind_title, years, isMailQC, couple=False, secondary_summary=None, secondary_ind_title=None, cfg=None):
    # Format years dynamically (e.g., "2021 & 2022" or "2021, 2022, ..., & 2024")
    if len(years) == 2:
        formatted_years = f"{years[0]} & {years[1]}"
    elif len(years) > 2:
        formatted_years = f"{', '.join(map(str, years[:-1]))}, & {years[-1]}"
    else:
        formatted_years = str(years[0])

    # Title: Centered, Dark Gray, 14pt, Calibri
    default_title = f'Summary of your {formatted_years} Tax Declarations'
    title_text = get_text(cfg, 'docTitle', default_title, yearRange=formatted_years)
    title = doc.add_paragraph(title_text)
    title_style = get_style(cfg, 'docTitle') or {'fontSize': 14, 'color': '#414141'}
    apply_alignment(title, title_style)
    if not get_style(cfg, 'docTitle'):
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    apply_style(run, title_style)

    # Blank line before name
    doc.add_paragraph()

    # Recipient Name for Primary Individual
    para = doc.add_paragraph()
    primary_first_name = primary_summary["tax_summary"]["first_name"]
    primary_last_name = primary_summary["tax_summary"]["last_name"]
    province = primary_summary["tax_summary"]["province"]
    para.add_run(f"{ind_title}. {primary_last_name}")

    # Display Secondary Individual's name if applicable
    if couple and secondary_summary and secondary_ind_title:
        secondary_first_name = secondary_summary["tax_summary"]["first_name"]
        secondary_last_name = secondary_summary["tax_summary"]["last_name"]
        para.add_run(f" & {secondary_ind_title}. {secondary_last_name}")

    # Blank line after name
    doc.add_paragraph()

    # Introduction paragraph
    para = doc.add_paragraph()
    default_attach = f'We have attached all the documents related to your {formatted_years} tax {"declarations" if couple else "declaration"}.'
    styled_run(para, cfg, 'introAttachment', default_attach, {'bold': False}, yearRange=formatted_years)
    para.add_run(' ')
    styled_run(para, cfg, 'introPassword', 'The password for the file named "COPY" is your 9-digit Social Insurance Number.', {'bold': True})
    para.add_run('\n')
    default_copy_desc = f'This {"set of documents" if couple else "document"} is a full copy of your tax return.'
    styled_run(para, cfg, 'introCopyDescription', default_copy_desc, {'bold': False})
    para.add_run(' ')
    styled_run(para, cfg, 'introCopyNoPrint', 'You do not need to print it or sign it;', {'bold': True})
    para.add_run(' ')
    styled_run(para, cfg, 'introCopyKeep', 'please keep it for your records and review it carefully to ensure everything is accurate and complete.', {'bold': False})

    # "Very Important" — blank line before it
    doc.add_paragraph()
    para = doc.add_paragraph()
    vi_text = get_text(cfg, 'veryImportantHeading', '** Very Important:')
    very_important_run = para.add_run(vi_text + '\n')
    apply_style(very_important_run, get_style(cfg, 'veryImportantHeading') or {'bold': True, 'underline': True, 'color': '#cd3350'})

    if province == "Québec":
        if isMailQC:
            # Special message if isMailQC is True
            fed_title = get_text(cfg, 'qcMailFedTitle', 'Regarding your Federal tax return:')
            federal_title_run = para.add_run(f'{fed_title}\n')
            apply_style(federal_title_run, get_style(cfg, 'qcMailFedTitle') or {'bold': True, 'underline': True})

            ns_before = get_text(cfg, 'qcMailFedNotSubmittedBefore', 'Please be advised that your Federal tax return ')
            r = para.add_run(ns_before)
            apply_style(r, get_style(cfg, 'qcMailFedNotSubmittedBefore') or {'bold': True})
            ns_underline = get_text(cfg, 'qcMailFedNotSubmittedUnderline', 'has not been submitted')
            r = para.add_run(ns_underline)
            apply_style(r, get_style(cfg, 'qcMailFedNotSubmittedUnderline') or {'bold': True, 'underline': True})
            ns_after = get_text(cfg, 'qcMailFedNotSubmittedAfter', ' to the government yet.')
            r = para.add_run(f'{ns_after}\n')
            apply_style(r, get_style(cfg, 'qcMailFedNotSubmittedAfter') or {'bold': True})

            auth_form_count = len(years) * (2 if couple else 1)
            auth_text = get_text(cfg, 'qcMailFedAuthForm',
                f'Attached are {auth_form_count} Authorization Forms. Please e-sign them (or print & sign) and e-mail them back to us as soon as possible so we can EFILE your return.')
            para.add_run(f'{auth_text}\n')

            styled_run(para, cfg, 'qcMailSignPartF', 'Please sign Part F.', {'italic': True, 'bold': True})
            para.add_run('\n\n')

            # Québec tax return
            qc_title = get_text(cfg, 'qcMailQCTitle', 'Regarding your Québec tax return:')
            quebec_title_run = para.add_run(f'{qc_title}\n')
            apply_style(quebec_title_run, get_style(cfg, 'qcMailQCTitle') or {'bold': True, 'underline': True})

            cannot_efile = get_text(cfg, 'qcMailQCCannotEfile', 'Please note that your Québec tax return cannot be transmitted via Efile.')
            para.add_run(f'{cannot_efile}\n')

            qc_docs = [f'"QC {year} - {primary_first_name} {primary_last_name}.pdf"' for year in years]
            if couple and secondary_summary:
                qc_docs += [f'"QC {year} - {secondary_first_name} {secondary_last_name}.pdf"' for year in years]
            print_text = get_text(cfg, 'qcMailQCPrint',
                f'For that reason, you need to print {", ".join(qc_docs)}, sign at the bottom of page ##, and mail them to the following address:')
            para.add_run(f'{print_text}\n')

            # Address in italics
            addr_text = get_text(cfg, 'qcAddress', 'Revenu Québec\nC. P. 2500, succursale Place-Desjardins\nMontréal (Québec) H5B 1A3')
            address = para.add_run(f'{addr_text}\n\n')
            apply_style(address, get_style(cfg, 'qcAddress') or {'italic': True})

            on_behalf = get_text(cfg, 'qcMailOnBehalf',
                '*If you would like us to mail the declaration on your behalf, please email us the signed declaration (e-signature that looks like your signature), and we will print and mail it by registered mail (with a tracking number) to QC Revenue. Please note that there will be an additional service fee of $25 plus Canada Post fees.')
            para.add_run(f'{on_behalf}\n')

        else:
            ns_before = get_text(cfg, 'qcEfileNotSubmittedBefore', f'Please be advised that your tax {"returns " if couple else "return "}')
            r = para.add_run(ns_before)
            apply_style(r, get_style(cfg, 'qcEfileNotSubmittedBefore') or {'bold': True})
            ns_underline = get_text(cfg, 'qcEfileNotSubmittedUnderline', f'{"have" if couple else "has"} not been submitted')
            r = para.add_run(ns_underline)
            apply_style(r, get_style(cfg, 'qcEfileNotSubmittedUnderline') or {'bold': True, 'underline': True})
            ns_after = get_text(cfg, 'qcEfileNotSubmittedAfter', ' to the government yet.')
            r = para.add_run(f'{ns_after}\n')
            apply_style(r, get_style(cfg, 'qcEfileNotSubmittedAfter') or {'bold': True})

            auth_form_count = len(years) * (4 if couple else 2)
            auth_text = get_text(cfg, 'qcEfileAuthForms',
                f'Attached are {auth_form_count} Authorization Forms. Please e-sign them (or print & sign) and e-mail them back to us as soon as possible so we can EFILE your returns.')
            para.add_run(f'{auth_text}\n')

            # Federal & Québec signature instructions
            styled_run(para, cfg, 'qcEfileSignFed', 'For the Federal Form, please sign Part F.', {'italic': True, 'bold': True})
            para.add_run('\n')

            styled_run(para, cfg, 'qcEfileSignQC', 'For the Quebec Form, please sign at the end of section 4.', {'italic': True, 'bold': True})
    else:
        # Non-Québec clients (no Québec-related details)
        auth_form_count = len(years) * (2 if couple else 1)
        ns_before = get_text(cfg, 'nonQcNotSubmittedBefore', f'Please be advised that your tax {"returns " if couple else "return "}')
        r = para.add_run(ns_before)
        apply_style(r, get_style(cfg, 'nonQcNotSubmittedBefore') or {'bold': True})
        ns_underline = get_text(cfg, 'nonQcNotSubmittedUnderline', f'{"have" if couple else "has"} not been submitted')
        r = para.add_run(ns_underline)
        apply_style(r, get_style(cfg, 'nonQcNotSubmittedUnderline') or {'bold': True, 'underline': True})
        ns_after = get_text(cfg, 'nonQcNotSubmittedAfter', ' to the government yet.')
        r = para.add_run(f'{ns_after}\n')
        apply_style(r, get_style(cfg, 'nonQcNotSubmittedAfter') or {'bold': True})

        auth_text = get_text(cfg, 'nonQcAuthForm',
            f'Attached are {auth_form_count} Authorization Forms. Please e-sign it (or print & sign) and e-mail it back to us as soon as possible so we can EFILE your return.')
        para.add_run(f'{auth_text}\n')

        styled_run(para, cfg, 'nonQcSignPartF', 'Please sign Part F.', {'italic': True, 'bold': True})
        para.add_run('\n')



def tax_return(doc, return_summary, year, last_year, secondary_summary=None, primary_title=None, secondary_title=None, isCouple=False, cfg=None):
    province = return_summary["tax_summary"]["province"]

    # Blank line above RESULTS
    doc.add_paragraph()
    para = doc.add_paragraph()
    results_text = get_text(cfg, 'resultsHeading', 'RESULTS')
    bold_red_results = para.add_run(f'{results_text}\n')
    apply_style(bold_red_results, get_style(cfg, 'resultsHeading') or {'bold': True, 'color': '#cd3350'})

    # Process primary individual's tax details
    if isCouple:
        para.add_run(f"\n{primary_title}. {return_summary['tax_summary']['last_name']}\n")

    # Federal Tax Return for primary individual
    fed_label = get_text(cfg, 'federalReturnLabel', 'Federal Tax return')
    r = para.add_run(f'{fed_label}\n')
    apply_style(r, get_style(cfg, 'federalReturnLabel') or {'bold': True})
    federal_refund = return_summary["tax_summary"]["federal_refund"]
    federal_owing = return_summary["tax_summary"]["federal_owing"]

    if federal_refund > 2:
        refund_prefix = get_text(cfg, 'refundPrefix', 'You are entitled to a refund of')
        para.add_run(f"{refund_prefix} ").bold = False
        refund_run = para.add_run(f"${federal_refund:,.2f} \n")
        refund_run.bold = True
        refund_run.font.color.rgb = RGBColor(0, 128, 0)
    elif federal_owing > 2:
        owing_prefix = get_text(cfg, 'owingPrefix', 'You owe the amount of')
        para.add_run(f"{owing_prefix} ").bold = False
        owing_run = para.add_run(f"${federal_owing:,.2f} \n")
        owing_run.bold = True
        owing_run.font.color.rgb = RGBColor(255, 0, 0)
    else:
        no_bal = get_text(cfg, 'noBalance', 'You have no Refund or Balance due.')
        para.add_run(f"{no_bal}\n")

    if province == "Quebec":
        # Quebec Tax Return for primary individual
        qc_label = get_text(cfg, 'quebecReturnLabel', 'Quebec Tax return')
        r = para.add_run(f'{qc_label}\n')
        apply_style(r, get_style(cfg, 'quebecReturnLabel') or {'bold': True})
        quebec_refund = return_summary["tax_summary"]["quebec_refund"]
        quebec_owing = return_summary["tax_summary"]["quebec_owing"]

        if quebec_refund > 2:
            refund_prefix = get_text(cfg, 'refundPrefix', 'You are entitled to a refund of')
            para.add_run(f"{refund_prefix} ").bold = False
            refund_run = para.add_run(f"${quebec_refund:,.2f}\n")
            refund_run.bold = True
            refund_run.font.color.rgb = RGBColor(0, 128, 0)
        elif quebec_owing > 2:
            owing_prefix = get_text(cfg, 'owingPrefix', 'You owe the amount of')
            para.add_run(f"{owing_prefix} ").bold = False
            owing_run = para.add_run(f"${quebec_owing:,.2f}\n")
            owing_run.bold = True
            owing_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            no_bal = get_text(cfg, 'noBalance', 'You have no Refund or Balance due.')
            para.add_run(f"{no_bal}\n")

    # Add payment section for primary individual only if this is the last year
    if year == last_year:
        add_payment_section(doc, federal_owing, quebec_owing if province == "Quebec" else 0, year, last_year, cfg=cfg)

    if isCouple and "carryforward_amounts" in return_summary and return_summary["carryforward_amounts"]:
        carryforward_amounts(doc, return_summary, cfg=cfg)

    # Add secondary individual if isCouple is True and secondary_summary is provided
    if isCouple and secondary_summary:
        para.add_run(f"\n{secondary_title}. {secondary_summary['tax_summary']['last_name']}\n")

        # Federal Tax Return for secondary individual
        fed_label = get_text(cfg, 'federalReturnLabel', 'Federal Tax return')
        r = para.add_run(f'{fed_label}\n')
        apply_style(r, get_style(cfg, 'federalReturnLabel') or {'bold': True})
        federal_refund = secondary_summary["tax_summary"]["federal_refund"]
        federal_owing = secondary_summary["tax_summary"]["federal_owing"]

        if federal_refund > 2:
            refund_prefix = get_text(cfg, 'refundPrefix', 'You are entitled to a refund of')
            para.add_run(f"{refund_prefix} ").bold = False
            refund_run = para.add_run(f"${federal_refund:,.2f}\n")
            refund_run.bold = True
            refund_run.font.color.rgb = RGBColor(0, 128, 0)
        elif federal_owing > 2:
            owing_prefix = get_text(cfg, 'owingPrefix', 'You owe the amount of')
            para.add_run(f"{owing_prefix} ").bold = False
            owing_run = para.add_run(f"${federal_owing:,.2f}\n")
            owing_run.bold = True
            owing_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            no_bal = get_text(cfg, 'noBalance', 'You have no Refund or Balance due.')
            para.add_run(f"{no_bal}\n")

        # Quebec Tax Return for secondary individual
        if province == "Quebec":
            qc_label = get_text(cfg, 'quebecReturnLabel', 'Quebec Tax return')
            r = para.add_run(f'{qc_label}\n')
            apply_style(r, get_style(cfg, 'quebecReturnLabel') or {'bold': True})
            quebec_refund = secondary_summary["tax_summary"]["quebec_refund"]
            quebec_owing = secondary_summary["tax_summary"]["quebec_owing"]

            if quebec_refund > 2:
                refund_prefix = get_text(cfg, 'refundPrefix', 'You are entitled to a refund of')
                para.add_run(f"{refund_prefix} ").bold = False
                refund_run = para.add_run(f"${quebec_refund:,.2f}")
                refund_run.bold = True
                refund_run.font.color.rgb = RGBColor(0, 128, 0)
            elif quebec_owing > 2:
                owing_prefix = get_text(cfg, 'owingPrefix', 'You owe the amount of')
                para.add_run(f"{owing_prefix} ").bold = False
                owing_run = para.add_run(f"${quebec_owing:,.2f}")
                owing_run.bold = True
                owing_run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                no_bal = get_text(cfg, 'noBalance', 'You have no Refund or Balance due.')
                para.add_run(f"{no_bal}\n")

        # Add payment section for secondary individual only if this is the last year
        add_payment_section(doc, federal_owing, quebec_owing if province == "Quebec" else 0, year, last_year, cfg=cfg)

        if "carryforward_amounts" in secondary_summary and secondary_summary["carryforward_amounts"]:
            carryforward_amounts(doc, secondary_summary, cfg=cfg)



# Function to add the payment section
def add_payment_section(doc, federal_owing, quebec_owing, year, last_year, cfg=None):


    if federal_owing > 2 or quebec_owing > 2:
        para = doc.add_paragraph()

        if federal_owing > 2 and quebec_owing > 2:
            owing_text = get_text(cfg, 'paymentOwingFedAndQC', 'You owe an amount on your Federal and Quebec returns;')
        elif federal_owing > 2:
            owing_text = get_text(cfg, 'paymentOwingFed', 'You owe an amount on your Federal return;')
        else:
            owing_text = get_text(cfg, 'paymentOwingQC', 'You owe an amount on your Quebec return;')
        r = para.add_run(f'{owing_text} ')
        apply_style(r, get_style(cfg, 'paymentOwingFed') or {'italic': True})

        if last_year == year:
            deadline = get_text(cfg, 'paymentDeadline', f'please make sure to pay the balance due by April 30, {int(year) + 1}, to avoid paying any interest.', dueYear=str(int(year) + 1))
            r = para.add_run(f'{deadline} ')
            apply_style(r, get_style(cfg, 'paymentDeadline') or {'italic': True})

        howto = get_text(cfg, 'paymentHowTo', 'Please wait a few days after we E-file to pay your outstanding balance. For more details on how to pay the amount due, please click on:')
        r = para.add_run(f'{howto} ')
        apply_style(r, get_style(cfg, 'paymentHowTo') or {'italic': True})

        # Add links for payment based on owing status
        if federal_owing > 2:
            add_hyperlink(para, 'Federal', 'https://www.canada.ca/en/revenue-agency/services/payments/payments-cra.html')

        if federal_owing > 2 and quebec_owing > 2:
            para.add_run(' and ')
        if quebec_owing > 2:
            add_hyperlink(para, 'Quebec', 'https://www.revenuquebec.ca/en/citizens/income-tax-return/paying-a-balance-due-or-receiving-a-refund/income-tax-balance-due/')


# Section: Solidarity Credit
def solidarity_credit(doc, return_summary, year, last_year, cfg=None):
    para = doc.add_paragraph()
    # Add title directly with bold and underline
    sol_title = get_text(cfg, 'solidarityTitle', 'Solidarity Credits:')
    title_run = para.add_run(f'{sol_title}\n')
    apply_style(title_run, get_style(cfg, 'solidarityTitle') or {'bold': True, 'underline': True})

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    solidarity_amount = return_summary["solidarity_amounts"]["solidarity_credit_amount"]

    para.add_run(f'You will receive a total of ').bold = False
    total_run = para.add_run(f"${solidarity_amount:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' for the ')
    para.add_run(f'Solidarity tax credit').bold = True
    if year == last_year:
        para.add_run(f' as follows:\n')
    else:
        para.add_run(f'.\n')

    # Extract amounts for each month
    amounts = {
        "July": return_summary["solidarity_amounts"]["july_amount"],
        "August": return_summary["solidarity_amounts"]["august_amount"],
        "September": return_summary["solidarity_amounts"]["september_amount"],
        "October": return_summary["solidarity_amounts"]["october_amount"],
        "November": return_summary["solidarity_amounts"]["november_amount"],
        "December": return_summary["solidarity_amounts"]["december_amount"],
        "January": return_summary["solidarity_amounts"]["january_amount"],
        "February": return_summary["solidarity_amounts"]["february_amount"],
        "March": return_summary["solidarity_amounts"]["march_amount"],
        "April": return_summary["solidarity_amounts"]["april_amount"],
        "May": return_summary["solidarity_amounts"]["may_amount"],
        "June": return_summary["solidarity_amounts"]["june_amount"],
    }

    # Filter out months with a zero amount
    filtered_months = {month: amount for month, amount in amounts.items() if amount > 0}

    if year == last_year:
        if len(filtered_months) <= 4:
            for month, amount in filtered_months.items():
                year = year_plus1 if month in ["July", "August", "September", "October", "November", "December"] else year_plus2
                para.add_run(f'{month} {year}: ${amount:,.2f}\n')
        else:
            # Check for consistency of amounts for the range
            unique_amounts = set(filtered_months.values())

            if len(unique_amounts) == 1:
                # If all amounts are the same
                amount = next(iter(unique_amounts))  # Get the single value
                para.add_run(f'${amount:,.2f}/month from July {year_plus1} to June {year_plus2}\n')
            else:
                # Group amounts by ranges
                amount_ranges = []
                current_amount = None
                start_month = None

                for month, amount in filtered_months.items():
                    year = year_plus1 if month in ["July", "August", "September", "October", "November", "December"] else year_plus2

                    if current_amount is None:
                        current_amount = amount
                        start_month = month
                    elif abs(current_amount - amount) > 0.3:  # Different enough to create a new range
                        end_year = year_plus1 if start_month in ["July", "August", "September", "October", "November", "December"] else year_plus2
                        amount_ranges.append((start_month, month, current_amount, end_year))
                        current_amount = amount
                        start_month = month
                    # Handle the last range
                    if month == "June":
                        end_year = year_plus2
                        amount_ranges.append((start_month, month, current_amount, end_year))

                for start, end, amount, end_year in amount_ranges:
                    para.add_run(f'${amount:,.2f}/month from {start} {year_plus1} to {end} {end_year}\n')


# Section: GST Credit
def gst_credit(doc, return_summary, isNewcomer, year, last_year, first_year, cfg=None):
    para = doc.add_paragraph()
    gst_title_text = get_text(cfg, 'gstTitle', 'GST/HST Credits:')
    title_run = para.add_run(gst_title_text)
    apply_style(title_run, get_style(cfg, 'gstTitle') or {'bold': True, 'underline': True})

    if isNewcomer and year == first_year:
        title_run = para.add_run(' *')
        title_run.bold = True
        title_run.underline = True
    para.add_run('\n')

    gst_credit_amount = return_summary["gst_amounts"]["gst_credit_amount"]
    july_amount = return_summary["gst_amounts"].get("july_amount", 0)
    october_amount = return_summary["gst_amounts"].get("october_amount", 0)
    january_amount = return_summary["gst_amounts"].get("january_amount", 0)
    april_amount = return_summary["gst_amounts"].get("april_amount", 0)

    para.add_run(f'You will receive a total of ').bold = False
    total_run = para.add_run(f"${gst_credit_amount:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' for the ')
    para.add_run(f'GST/ HST credit').bold = True
    if year == last_year:
        para.add_run(f' as follows:\n')
    else:
        para.add_run(f'.\n')

    if year == last_year:
        if july_amount > 0:
            para.add_run(f'July {year}: ${july_amount:,.2f}\n')
        if october_amount > 0:
            para.add_run(f'October {year}: ${october_amount:,.2f}\n')
        if january_amount > 0:
            para.add_run(f'January {year + 1}: ${january_amount:,.2f}\n')
        if april_amount > 0:
            para.add_run(f'April {year + 1}: ${april_amount:,.2f}\n')

    if isNewcomer and year == first_year:
        para.add_run(
            '\n*Note that you will receive a letter from Canada Revenue Agency asking you to provide your income before arrival to Canada (so from January 1st until the date of arrival). '
            'Even though it was mentioned on the declaration, you still need to respond to the letter and provide the amount. If you do not reply, they will not pay the GST amount.\n'
        ).italic = True


# Section: ECGEB Credit
def ecgeb_credit(doc, return_summary, isNewcomer, year, last_year, first_year, cfg=None):
    para = doc.add_paragraph()
    ecgeb_title_text = get_text(cfg, 'ecgebTitle', 'Canada Groceries & Essentials Benefit:')
    title_run = para.add_run(ecgeb_title_text)
    apply_style(title_run, get_style(cfg, 'ecgebTitle') or {'bold': True, 'underline': True})

    if isNewcomer and year == first_year:
        title_run = para.add_run(' *')
        title_run.bold = True
        title_run.underline = True
    para.add_run('\n')

    ecgeb_credit_amount = return_summary["ecgeb_amounts"]["ecgeb_credit_amount"]
    july_amount = return_summary["ecgeb_amounts"].get("july_amount", 0)
    october_amount = return_summary["ecgeb_amounts"].get("october_amount", 0)
    january_amount = return_summary["ecgeb_amounts"].get("january_amount", 0)
    april_amount = return_summary["ecgeb_amounts"].get("april_amount", 0)

    para.add_run(f'You will receive a total of ').bold = False
    total_run = para.add_run(f"${ecgeb_credit_amount:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' for the ')
    para.add_run(f'Canada Groceries and Essentials Benefit (ECGEB)').bold = True
    if year == last_year:
        para.add_run(f' as follows:\n')
    else:
        para.add_run(f'.\n')

    if year == last_year:
        if july_amount > 0:
            para.add_run(f'July {year}: ${july_amount:,.2f}\n')
        if october_amount > 0:
            para.add_run(f'October {year}: ${october_amount:,.2f}\n')
        if january_amount > 0:
            para.add_run(f'January {year + 1}: ${january_amount:,.2f}\n')
        if april_amount > 0:
            para.add_run(f'April {year + 1}: ${april_amount:,.2f}\n')

    if isNewcomer and year == first_year:
        para.add_run(
            '\n*Note that you will receive a letter from Canada Revenue Agency asking you to provide your income before arrival to Canada (so from January 1st until the date of arrival). '
            'Even though it was mentioned on the declaration, you still need to respond to the letter and provide the amount. If you do not reply, they will not pay the ECGEB amount.\n'
        ).italic = True


# Section: Summary of Carryforward Amounts
def carryforward_amounts(doc, return_summary, cfg=None):
    para = doc.add_paragraph()
    tuition_title = get_text(cfg, 'tuitionTitle', 'Your accumulated tuition fees carried forward to future years:')
    title_run = para.add_run(f'{tuition_title}\n')
    apply_style(title_run, get_style(cfg, 'tuitionTitle') or {'bold': True})

    federal_tuition_amount = return_summary["carryforward_amounts"]["federal_tuition_amount"]
    quebec_tuition_8_percent = return_summary["carryforward_amounts"]["quebec_tuition_8_percent"]

    # Convert string values safely (removing commas for conversion)
    if isinstance(federal_tuition_amount, str):
        if federal_tuition_amount.strip():
            federal_tuition_amount = int(federal_tuition_amount.replace(",", ""))
        else:
            federal_tuition_amount = 0
    if isinstance(quebec_tuition_8_percent, str):
        if quebec_tuition_8_percent.strip():
            quebec_tuition_8_percent = int(quebec_tuition_8_percent.replace(",", ""))
        else:
            quebec_tuition_8_percent = 0

    # Format numbers with commas as thousand separators when adding them to the document
    if federal_tuition_amount > 0:
        fed_label = get_text(cfg, 'tuitionFedLabel', 'Federal (eligible to 15%):  $')
        para.add_run(fed_label)
        para.add_run(f"{federal_tuition_amount:,}")
    if quebec_tuition_8_percent > 0:
        qc_label = get_text(cfg, 'tuitionQCLabel', 'QC (eligible to 8%): $')
        para.add_run(f'\n{qc_label}')
        para.add_run(f"{quebec_tuition_8_percent:,}")

    tuition_exp = get_text(cfg, 'tuitionExplanation',
        'Those accumulated tuition fees are tax credits that you will be using in future tax declarations when you work and pay tax on your income.')
    r = para.add_run(f'\n\n{tuition_exp}')
    apply_style(r, get_style(cfg, 'tuitionExplanation') or {'italic': True})



# Section: Child Benefit Amount
def child_benefit(doc, return_summary, year, last_year, cfg=None):
    para = doc.add_paragraph()
    # Add title directly with bold and underline
    ccb_title = get_text(cfg, 'ccbTitle', 'Canada Child Benefit (CCB):')
    title_run = para.add_run(f'{ccb_title}\n')
    apply_style(title_run, get_style(cfg, 'ccbTitle') or {'bold': True, 'underline': True})

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    # Extract total Canada Child Benefit amount
    total_ccb = return_summary["ccb_amounts"]["ccb_amount"]
    amounts = {
        f"July {year_plus1}": return_summary["ccb_amounts"]["july_amount"],
        f"August {year_plus1}": return_summary["ccb_amounts"]["august_amount"],
        f"September {year_plus1}": return_summary["ccb_amounts"]["september_amount"],
        f"October {year_plus1}": return_summary["ccb_amounts"]["october_amount"],
        f"November {year_plus1}": return_summary["ccb_amounts"]["november_amount"],
        f"December {year_plus1}": return_summary["ccb_amounts"]["december_amount"],
        f"January {year_plus2}": return_summary["ccb_amounts"]["january_amount"],
        f"February {year_plus2}": return_summary["ccb_amounts"]["february_amount"],
        f"March {year_plus2}": return_summary["ccb_amounts"]["march_amount"],
        f"April {year_plus2}": return_summary["ccb_amounts"]["april_amount"],
        f"May {year_plus2}": return_summary["ccb_amounts"]["may_amount"],
        f"June {year_plus2}": return_summary["ccb_amounts"]["june_amount"],
    }

    # Display total benefit amount
    para.add_run(f'You will receive a total of ').bold = False
    total_run = para.add_run(f"${total_ccb:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' for the ')
    para.add_run(f'Canada Child Benefit ').bold = True
    if year == last_year:
        para.add_run(f' as follows:\n')
    else:
        para.add_run(f'.\n')

    # Filter out months with a zero amount
    filtered_months = {month: amount for month, amount in amounts.items() if amount > 0}

    if year == last_year:
        if len(filtered_months) <= 4:
            for month, amount in filtered_months.items():
                para.add_run(f'{month}: ${amount:,.2f}\n')
        else:
            # Check for consistency of amounts for the range
            unique_amounts = set(filtered_months.values())

            if len(unique_amounts) == 1:
                # If all amounts are the same
                amount = next(iter(unique_amounts))  # Get the single value
                para.add_run(f'${amount:,.2f}/month from July {year_plus1} to June {year_plus2}')
            else:
                # Group amounts by ranges
                amount_ranges = []
                current_amount = None
                start_month = None

                for month, amount in filtered_months.items():
                    # Automatically determine the year from month
                    year = month.split()[1]  # Extract year from month string

                    if current_amount is None:
                        current_amount = amount
                        start_month = month
                    elif abs(current_amount - amount) > 0.3:
                        amount_ranges.append((start_month, month, current_amount, year))
                        current_amount = amount
                        start_month = month

                # Handle the last range
                if start_month:
                    amount_ranges.append((start_month, month, current_amount, year))

                for start, end, amount, year in amount_ranges:
                    para.add_run(f'${amount:,.2f}/month from {start} to {end}')

# Section: Quebec Family Allowance
def quebec_family_allowance(doc, return_summary, year, last_year, cfg=None):
    para = doc.add_paragraph()

    fa_title = get_text(cfg, 'familyAllowanceTitle', 'Quebec Family Allowance:')
    title_run = para.add_run(fa_title + '\n')
    apply_style(title_run, get_style(cfg, 'familyAllowanceTitle') or {'bold': True, 'underline': True})

    total_allowance_amount = return_summary["family_allowance_amounts"]["fa_amount"]

    para.add_run(f'You will receive a total of ').bold = False
    total_run = para.add_run(f"${total_allowance_amount:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' for the ')
    fa_title = get_text(cfg, 'familyAllowanceTitle', 'Quebec Family Allowance:')
    para.add_run(fa_title).bold = True
    if year == last_year:
        para.add_run(f' as follows:\n')
    else:
        para.add_run(f'.\n')

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    july_amount = return_summary["family_allowance_amounts"]["july_amount"]
    october_amount = return_summary["family_allowance_amounts"]["october_amount"]
    january_amount = return_summary["family_allowance_amounts"]["january_amount"]
    april_amount = return_summary["family_allowance_amounts"]["april_amount"]

    if year == last_year:
        if july_amount > 0:
            para.add_run(f'July {year_plus1}: ${july_amount:,.2f}\n')
        if october_amount > 0:
            para.add_run(f'October {year_plus1}: ${october_amount:,.2f}\n')
        if january_amount > 0:
            para.add_run(f'January {year_plus2}: ${january_amount:,.2f}\n')
        if april_amount > 0:
            para.add_run(f'April {year_plus2}: ${april_amount:,.2f}\n')

# Section: Conclusion
def conclusion(doc, isMailQC=False, cfg=None):
    para = doc.add_paragraph()

    # Add the red text above the conclusion
    # Use singular "form" for Efile fed mail QC (Individual), plural "forms" otherwise
    form_text = "form" if isMailQC else "forms"
    waiting_default = f'We will be waiting for the signed authorization {form_text} to proceed.'
    waiting_text = get_text(cfg, 'conclusionWaiting', waiting_default)
    red_run = para.add_run(f'\n{waiting_text}\n')
    apply_style(red_run, get_style(cfg, 'conclusionWaiting') or {'color': '#cd3350'})

    # Add normal text for "Thank you."
    thank_text = get_text(cfg, 'thankYou', 'Thank you.')
    para.add_run(f'\n{thank_text}\n')

    # Add the main conclusion text with specified formatting
    disclaimer_default = (
        'We at Sankari Inc. are pleased to respond to your tax inquiries and/or file your tax returns based on the '
        'information that you provide. Inaccurate or incomplete information provided by you may lead to inadequate or '
        'incorrect advice for which Sankari Inc. team cannot be held responsible. You, the client, are responsible for '
        'giving correct information and documentation to Sankari Inc.'
    )
    disclaimer_text = get_text(cfg, 'disclaimer', disclaimer_default)
    conclusion_run = para.add_run(f'\n{disclaimer_text}')
    apply_style(conclusion_run, get_style(cfg, 'disclaimer') or {'fontSize': 8, 'color': '#7f7f7f', 'italic': True})


def createIndividualWordDocFR(individual, output_file_path, doc_text_config=None):
    doc = Document()
    set_default_font(doc, "Calibri", 10)

    cfg = get_cfg(doc_text_config, resolve_doc_type_key('FR', is_couple=False, is_multiyear=True))

    ind_title = individual[0]['title']
    if ind_title == 'Mr':
        ind_title = "M."
    elif ind_title in ['Ms', 'Mrs']:
        ind_title = "Mme"

    isMailQC = any(person.get('isMailQC', False) for person in individual)
    isNewcomer = any(person.get('isNewcomer', False) for person in individual)

    summaries = {person['year']: person['summary'] for person in individual}
    years = sorted(summaries.keys(), key=int)

    # Add section_1FR with formatted years
    section_1FR(doc, summaries[years[0]], ind_title, years, isMailQC, cfg=cfg)

    for year in years:
        return_summary = summaries[year]

        para = doc.add_paragraph()
        run = para.add_run(f"{year}")
        run.underline = True

        tax_returnFR(doc, return_summary, year, years[-1], cfg=cfg)

        if "gst_amounts" in return_summary and return_summary["gst_amounts"]:
            gst_creditFR(doc, return_summary, isNewcomer, year, years[-1], years[0], cfg=cfg)

        if "ecgeb_amounts" in return_summary and return_summary["ecgeb_amounts"]:
            ecgeb_creditFR(doc, return_summary, isNewcomer, year, years[-1], years[0], cfg=cfg)

        if "solidarity_amounts" in return_summary and return_summary["solidarity_amounts"]:
            solidarity_creditFR(doc, return_summary, year, years[-1], cfg=cfg)

        if "ccb_amounts" in return_summary and return_summary["ccb_amounts"]:
            child_benefitFR(doc, return_summary, year, years[-1], cfg=cfg)

        if "family_allowance_amounts" in return_summary and return_summary["family_allowance_amounts"]:
            quebec_family_allowanceFR(doc, return_summary, year, years[-1], cfg=cfg)

        if year == years[-1] and "carryforward_amounts" in return_summary and return_summary["carryforward_amounts"]:
            carryforward_amountsFR(doc, return_summary, cfg=cfg)

    conclusionFR(doc, cfg=cfg)
    doc.save(output_file_path)


def section_1FR(doc, return_summary, ind_title, years, isMailQC, couple=False, secondary_summary=None, secondary_ind_title=None, cfg=None):
    # Format years dynamically
    if len(years) == 2:
        formatted_years = f"{years[0]} & {years[1]}"
    elif len(years) > 2:
        formatted_years = f"{', '.join(map(str, years[:-1]))}, & {years[-1]}"
    else:
        formatted_years = str(years[0])

    # Title: Centered, Dark Gray, 14pt, Calibri
    default_title = f"Sommaire de {'vos' if couple else 'votre'} d\u00e9clarations d\u2019imp\u00f4ts {formatted_years}"
    title_text = get_text(cfg, 'docTitle', default_title, yearRange=formatted_years)
    title = doc.add_paragraph(title_text)
    title_style = get_style(cfg, 'docTitle') or {'fontSize': 14, 'color': '#414141'}
    apply_alignment(title, title_style)
    if not get_style(cfg, 'docTitle'):
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    apply_style(run, title_style)

    # Blank line before name
    doc.add_paragraph()

    # Recipient Name
    para = doc.add_paragraph()
    first_name = return_summary["tax_summary"]["first_name"]
    last_name = return_summary["tax_summary"]["last_name"]
    province = return_summary["tax_summary"]["province"]
    para.add_run(f"{ind_title} {last_name}")

    # Display Secondary Individual's name if applicable
    if couple and secondary_summary and secondary_ind_title:
        secondary_first_name = secondary_summary["tax_summary"]["first_name"]
        secondary_last_name = secondary_summary["tax_summary"]["last_name"]
        para.add_run(f" & {secondary_ind_title} {secondary_last_name}")

    # Blank line after name
    doc.add_paragraph()

    # Introduction paragraph
    para = doc.add_paragraph()
    default_attach = f'Nous avons joint \u00e0 ce courriel tous les documents de {"\u0076os d\u00e9clarations" if couple else "votre d\u00e9claration"} d\u2019imp\u00f4ts {formatted_years}.'
    styled_run(para, cfg, 'introAttachment', default_attach, {'bold': False}, yearRange=formatted_years)
    para.add_run(' ')
    styled_run(para, cfg, 'introPassword', 'Le mot de passe du document nomm\u00e9 \u201cCOPIE\u201d est compos\u00e9 des neuf chiffres de votre num\u00e9ro d\u2019assurance sociale.', {'bold': True})
    para.add_run('\n')
    if couple:
        default_copy_desc = 'Ce jeu de documents constitue une copie compl\u00e8te de vos d\u00e9clarations de revenus.'
        default_copy_noprint = 'Vous n\'avez pas besoin de les imprimer ou de les signer;'
        default_copy_keep = 'vous avez juste besoin de les retenir pour votre dossier. Veuillez revoir les d\u00e9clarations de revenu attentivement afin de vous assurer qu\u2019elles sont exactes et compl\u00e8tes.'
    else:
        default_copy_desc = 'Ce document constitue une copie compl\u00e8te de votre d\u00e9claration de revenus.'
        default_copy_noprint = 'Vous n\'avez pas besoin de l\u2019imprimer ou de le signer;'
        default_copy_keep = 'vous avez juste besoin de le retenir pour votre dossier. Veuillez revoir la d\u00e9claration de revenu attentivement afin de vous assurer qu\u2019elle est exacte et compl\u00e8te.'
    styled_run(para, cfg, 'introCopyDescription', default_copy_desc, {'bold': False})
    para.add_run(' ')
    styled_run(para, cfg, 'introCopyNoPrint', default_copy_noprint, {'bold': True})
    styled_run(para, cfg, 'introCopyKeep', default_copy_keep, {'bold': False})

    # "Very Important" — blank line before it
    doc.add_paragraph()
    para = doc.add_paragraph()
    vi_text = get_text(cfg, 'veryImportantHeading', '** TR\u00c8S IMPORTANT :')
    very_important_run = para.add_run(vi_text + '\n')
    apply_style(very_important_run, get_style(cfg, 'veryImportantHeading') or {'bold': True, 'underline': True, 'color': '#cd3350'})

    if province == "Qu\u00e9bec":
        if isMailQC:
            # Special message if isMailQC is True
            fed_title = get_text(cfg, 'qcMailFedTitle', 'D\u00e9claration F\u00e9d\u00e9rale :')
            federal_title_run = para.add_run(f'{fed_title}\n')
            apply_style(federal_title_run, get_style(cfg, 'qcMailFedTitle') or {'bold': True, 'underline': True})

            ns_before = get_text(cfg, 'qcMailFedNotSubmittedBefore', 'Veuillez noter que votre d\u00e9claration f\u00e9d\u00e9rale ')
            r = para.add_run(ns_before)
            apply_style(r, get_style(cfg, 'qcMailFedNotSubmittedBefore') or {'bold': True})
            ns_underline = get_text(cfg, 'qcMailFedNotSubmittedUnderline', 'n\u2019a pas encore \u00e9t\u00e9 soumise')
            r = para.add_run(ns_underline)
            apply_style(r, get_style(cfg, 'qcMailFedNotSubmittedUnderline') or {'bold': True, 'underline': True})
            ns_after = get_text(cfg, 'qcMailFedNotSubmittedAfter', ' au gouvernement.')
            r = para.add_run(f'{ns_after}\n')
            apply_style(r, get_style(cfg, 'qcMailFedNotSubmittedAfter') or {'bold': True})

            auth_form_count = len(years) * (2 if couple else 1)
            auth_text = get_text(cfg, 'qcMailFedAuthForm',
                f'Vous trouverez ci-joint {auth_form_count} formulaire(s) d\u2019autorisation. Veuillez les signer \u00e9lectroniquement (ou les imprimer et signer) et nous les envoyer par courriel d\u00e8s que possible afin que nous puissions transmettre votre d\u00e9claration par EFILE.')
            para.add_run(f'{auth_text}\n')

            styled_run(para, cfg, 'qcMailSignPartF', 'S\u2019il vous pla\u00eet signer la partie F.', {'italic': True, 'bold': True})
            para.add_run('\n\n')

            qc_docs = [f'"QC {year} - {first_name} {last_name}.pdf"' for year in years]

            # Québec tax return
            qc_title = get_text(cfg, 'qcMailQCTitle', 'D\u00e9claration Provinciale :')
            quebec_title_run = para.add_run(f'{qc_title}\n')
            apply_style(quebec_title_run, get_style(cfg, 'qcMailQCTitle') or {'bold': True, 'underline': True})

            cannot_efile = get_text(cfg, 'qcMailQCCannotEfile', 'Veuillez noter que votre d\u00e9claration provinciale ne peut pas \u00eatre transmise via Efile.')
            para.add_run(f'{cannot_efile}\n')

            if couple and secondary_summary:
                qc_docs += [f'"QC {year} - {secondary_first_name} {secondary_last_name}.pdf"' for year in years]
            print_text = get_text(cfg, 'qcMailQCPrint',
                f'Vous devez \u00e9galement imprimer {", ".join(qc_docs)}, signer en bas de la page ##, et l\u2019envoyer par la poste \u00e0 l\u2019adresse suivante :')
            para.add_run(f'{print_text}\n')

            addr_text = get_text(cfg, 'qcAddress', 'Revenu Qu\u00e9bec\nC. P. 2500, succursale Place-Desjardins\nMontr\u00e9al (Qu\u00e9bec) H5B 1A3')
            address = para.add_run(f'{addr_text}\n')
            apply_style(address, get_style(cfg, 'qcAddress') or {'italic': True})

            on_behalf = get_text(cfg, 'qcMailOnBehalf',
                '*Si vous souhaitez que nous envoyions la d\u00e9claration pour vous, veuillez nous envoyer la d\u00e9claration sign\u00e9e (signature \u00e9lectronique ressemblant \u00e0 votre signature), et nous l\u2019imprimerons et l\u2019enverrons par courrier recommand\u00e9 avec un num\u00e9ro de suivi \u00e0 Revenu Qu\u00e9bec. Des frais de service de 25 $ plus les frais de Postes Canada s\u2019appliquent.')
            para.add_run(f'{on_behalf}\n')

        else:
            ns_before = get_text(cfg, 'qcEfileNotSubmittedBefore', f'{"Veuillez noter que vos d\u00e9clarations d\u2019imp\u00f4t " if couple else "Veuillez noter que votre d\u00e9claration d\u2019imp\u00f4t "}')
            r = para.add_run(ns_before)
            apply_style(r, get_style(cfg, 'qcEfileNotSubmittedBefore') or {'bold': True})
            ns_underline = get_text(cfg, 'qcEfileNotSubmittedUnderline', f'{"n\u2019ont pas encore \u00e9t\u00e9 soumises" if couple else "n\u2019a pas encore \u00e9t\u00e9 soumise"}')
            r = para.add_run(ns_underline)
            apply_style(r, get_style(cfg, 'qcEfileNotSubmittedUnderline') or {'bold': True, 'underline': True})
            ns_after = get_text(cfg, 'qcEfileNotSubmittedAfter', ' au gouvernement.')
            r = para.add_run(f'{ns_after}\n')
            apply_style(r, get_style(cfg, 'qcEfileNotSubmittedAfter') or {'bold': True})

            auth_form_count = len(years) * (4 if couple else 2)
            auth_text = get_text(cfg, 'qcEfileAuthForms',
                f'Vous trouverez ci-joint {auth_form_count} formulaire(s) d\u2019autorisation. Veuillez les signer \u00e9lectroniquement (ou les imprimer et signer) et nous les envoyer par courriel d\u00e8s que possible afin que nous puissions transmettre votre d\u00e9claration par EFILE.')
            para.add_run(f'{auth_text}\n')

            styled_run(para, cfg, 'qcEfileSignFed', 'Pour le formulaire f\u00e9d\u00e9ral, veuillez signer la partie F.', {'italic': True, 'bold': True})
            para.add_run('\n')

            styled_run(para, cfg, 'qcEfileSignQC', 'Pour le formulaire du Qu\u00e9bec, veuillez signer \u00e0 la fin de la section 4.', {'italic': True, 'bold': True})
    else:
        auth_form_count = len(years) * (2 if couple else 1)
        # Original message if isMailQC is False
        ns_before = get_text(cfg, 'nonQcNotSubmittedBefore', f'{"Noter que vos d\u00e9clarations d\u2019imp\u00f4t " if couple else "Noter que votre d\u00e9claration d\u2019imp\u00f4t "}')
        r = para.add_run(ns_before)
        apply_style(r, get_style(cfg, 'nonQcNotSubmittedBefore') or {'bold': True})
        ns_underline = get_text(cfg, 'nonQcNotSubmittedUnderline', f'{"n\u2019ont pas encore \u00e9t\u00e9 soumises" if couple else "n\u2019a pas encore \u00e9t\u00e9 soumise"}')
        r = para.add_run(ns_underline)
        apply_style(r, get_style(cfg, 'nonQcNotSubmittedUnderline') or {'bold': True, 'underline': True})
        ns_after = get_text(cfg, 'nonQcNotSubmittedAfter', ' au gouvernement via EFile.')
        r = para.add_run(f'{ns_after}\n')
        apply_style(r, get_style(cfg, 'nonQcNotSubmittedAfter') or {'bold': True})

        auth_text = get_text(cfg, 'nonQcAuthForm',
            f'Vous trouverez ci-joint {auth_form_count} formulaires d\u2019autorisation. Veuillez les signer \u00e9lectroniquement (ou imprimer/signer) et nous les envoyer par courriel le plus t\u00f4t possible afin que nous puissions transmettre votre d\u00e9claration par EFILE.')
        para.add_run(f'{auth_text}\n')

        styled_run(para, cfg, 'nonQcSignPartF', 'Veuillez signer la partie F.', {'italic': True, 'bold': True})
        para.add_run('\n')

def tax_returnFR(doc, return_summary, year, last_year, secondary_summary=None, primary_title=None, secondary_title=None, isCouple=False, cfg=None):
    province = return_summary["tax_summary"]["province"]

    # Blank line above RESULTS
    doc.add_paragraph()
    para = doc.add_paragraph()

    # Title "RÉSULTATS" in bold red
    results_text = get_text(cfg, 'resultsHeading', 'R\u00c9SULTATS')
    bold_red_results = para.add_run(f'{results_text}\n')
    apply_style(bold_red_results, get_style(cfg, 'resultsHeading') or {'bold': True, 'color': '#cd3350'})

    # Helper function to format numbers in the French style
    def format_french_number(amount):
        return f"{amount:,.2f}".replace(",", " ").replace(".", ",")

        # Function to handle payment section if amounts are owed
    def add_payment_section(para, federal_owing, quebec_owing, year, last_year):
        if federal_owing > 0 or quebec_owing > 0:
            para = doc.add_paragraph()

            if federal_owing > 0 and quebec_owing > 0:
                owing_text = get_text(cfg, 'paymentOwingFedAndQC', 'Vous devez un montant \u00e0 votre d\u00e9claration f\u00e9d\u00e9rale et provinciale;')
            elif federal_owing > 0:
                owing_text = get_text(cfg, 'paymentOwingFed', 'Vous devez un montant \u00e0 votre d\u00e9claration f\u00e9d\u00e9rale;')
            else:
                owing_text = get_text(cfg, 'paymentOwingQC', 'Vous devez un montant \u00e0 votre d\u00e9claration provinciale;')
            r = para.add_run(f'{owing_text} ')
            apply_style(r, get_style(cfg, 'paymentOwingFed') or {'italic': True})

            if last_year == year:
                deadline = get_text(cfg, 'paymentDeadline', f'assurez-vous de payer le solde d\u00fb avant le 30 avril {int(year) + 1} pour \u00e9viter de payer des int\u00e9r\u00eats.', dueYear=str(int(year) + 1))
                r = para.add_run(f'{deadline} ')
                apply_style(r, get_style(cfg, 'paymentDeadline') or {'italic': True})

            howto = get_text(cfg, 'paymentHowTo', 'Veuillez attendre quelques jours apr\u00e8s la transmission Efile pour payer votre solde d\u00fb. Pour plus de d\u00e9tails sur la fa\u00e7on de payer le montant d\u00fb, veuillez cliquer sur :')
            r = para.add_run(f'{howto} ')
            apply_style(r, get_style(cfg, 'paymentHowTo') or {'italic': True})

            # Add hyperlinks for payment details
            if federal_owing > 0:
                add_hyperlink(para, 'F\u00e9d\u00e9ral', 'https://www.canada.ca/fr/agence-revenu/services/paiements/paiements-arc.html')
            if federal_owing > 0 and quebec_owing > 0:
                para.add_run(' et ')
            if quebec_owing > 0:
                add_hyperlink(para, 'Qu\u00e9bec', 'https://www.revenuquebec.ca/fr/citoyens/declaration-de-revenus/payer-ou-etre-rembourse/solde-dimpot-a-payer/')


    # Process primary individual's tax details
    if isCouple:
        para.add_run(f"\n{primary_title} {return_summary['tax_summary']['last_name']}\n")

    # Federal Tax Return for primary individual
    fed_label = get_text(cfg, 'federalReturnLabel', 'D\u00e9claration F\u00e9d\u00e9rale')
    r = para.add_run(f'{fed_label}\n')
    apply_style(r, get_style(cfg, 'federalReturnLabel') or {'bold': True})
    federal_refund = return_summary["tax_summary"]["federal_refund"]
    federal_owing = return_summary["tax_summary"]["federal_owing"]

    if federal_refund > 0:
        refund_prefix = get_text(cfg, 'refundPrefix', 'Vous avez droit \u00e0 un remboursement de')
        para.add_run(f"{refund_prefix} ").bold = False
        refund_run = para.add_run(f"{format_french_number(federal_refund)} $ \n")
        refund_run.bold = True
        refund_run.font.color.rgb = RGBColor(0, 128, 0)
    elif federal_owing > 0:
        owing_prefix = get_text(cfg, 'owingPrefix', 'Vous devez un montant de')
        para.add_run(f"{owing_prefix} ").bold = False
        owing_run = para.add_run(f"{format_french_number(federal_owing)} $ \n")
        owing_run.bold = True
        owing_run.font.color.rgb = RGBColor(255, 0, 0)
    else:
        no_bal = get_text(cfg, 'noBalance', 'Vous n\u2019avez ni remboursement ni solde d\u00fb.')
        para.add_run(f"{no_bal}\n")

    if province == "Qu\u00e9bec":
        # Quebec Tax Return for primary individual
        qc_label = get_text(cfg, 'quebecReturnLabel', 'D\u00e9claration Provinciale')
        r = para.add_run(f'{qc_label}\n')
        apply_style(r, get_style(cfg, 'quebecReturnLabel') or {'bold': True})
        quebec_refund = return_summary["tax_summary"]["quebec_refund"]
        quebec_owing = return_summary["tax_summary"]["quebec_owing"]

        if quebec_refund > 0:
            refund_prefix = get_text(cfg, 'refundPrefix', 'Vous avez droit \u00e0 un remboursement de')
            para.add_run(f"{refund_prefix} ").bold = False
            refund_run = para.add_run(f"{format_french_number(quebec_refund)} $ \n")
            refund_run.bold = True
            refund_run.font.color.rgb = RGBColor(0, 128, 0)
        elif quebec_owing > 0:
            owing_prefix = get_text(cfg, 'owingPrefix', 'Vous devez un montant de')
            para.add_run(f"{owing_prefix} ").bold = False
            owing_run = para.add_run(f"{format_french_number(quebec_owing)} $ \n")
            owing_run.bold = True
            owing_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            no_bal = get_text(cfg, 'noBalance', 'Vous n\u2019avez ni remboursement ni solde d\u00fb.')
            para.add_run(f"{no_bal}\n")

    # Add payment section for primary individual if this is the last year
    if year == last_year:
        add_payment_section(doc, federal_owing, quebec_owing if province == "Qu\u00e9bec" else 0, year, last_year)

    if isCouple and "carryforward_amounts" in return_summary and return_summary["carryforward_amounts"]:
        carryforward_amounts(doc, return_summary, cfg=cfg)

    # Process secondary individual if applicable
    if isCouple and secondary_summary:
        para.add_run(f"\n{secondary_title} {secondary_summary['tax_summary']['last_name']}\n")

        # Federal Tax Return for secondary individual
        fed_label = get_text(cfg, 'federalReturnLabel', 'D\u00e9claration F\u00e9d\u00e9rale')
        r = para.add_run(f'{fed_label}\n')
        apply_style(r, get_style(cfg, 'federalReturnLabel') or {'bold': True})
        federal_refund = secondary_summary["tax_summary"]["federal_refund"]
        federal_owing = secondary_summary["tax_summary"]["federal_owing"]

        if federal_refund > 0:
            refund_prefix = get_text(cfg, 'refundPrefix', 'Vous avez droit \u00e0 un remboursement de')
            para.add_run(f"{refund_prefix} ").bold = False
            refund_run = para.add_run(f"{format_french_number(federal_refund)} $ \n")
            refund_run.bold = True
            refund_run.font.color.rgb = RGBColor(0, 128, 0)
        elif federal_owing > 0:
            owing_prefix = get_text(cfg, 'owingPrefix', 'Vous devez un montant de')
            para.add_run(f"{owing_prefix} ").bold = False
            owing_run = para.add_run(f"{format_french_number(federal_owing)} $ \n")
            owing_run.bold = True
            owing_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            no_bal = get_text(cfg, 'noBalance', 'Vous n\u2019avez ni remboursement ni solde d\u00fb.')
            para.add_run(f"{no_bal}\n")

    if province == "Qu\u00e9bec":
        # Quebec Tax Return for secondary individual
        qc_label = get_text(cfg, 'quebecReturnLabel', 'D\u00e9claration Provinciale')
        r = para.add_run(f'{qc_label}\n')
        apply_style(r, get_style(cfg, 'quebecReturnLabel') or {'bold': True})
        quebec_refund = secondary_summary["tax_summary"]["quebec_refund"]
        quebec_owing = secondary_summary["tax_summary"]["quebec_owing"]

        if quebec_refund > 0:
            refund_prefix = get_text(cfg, 'refundPrefix', 'Vous avez droit \u00e0 un remboursement de')
            para.add_run(f"{refund_prefix} ").bold = False
            refund_run = para.add_run(f"{format_french_number(quebec_refund)} $ \n")
            refund_run.bold = True
            refund_run.font.color.rgb = RGBColor(0, 128, 0)
        elif quebec_owing > 0:
            owing_prefix = get_text(cfg, 'owingPrefix', 'Vous devez un montant de')
            para.add_run(f"{owing_prefix} ").bold = False
            owing_run = para.add_run(f"{format_french_number(quebec_owing)} $ \n")
            owing_run.bold = True
            owing_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            no_bal = get_text(cfg, 'noBalance', 'Vous n\u2019avez ni remboursement ni solde d\u00fb.')
            para.add_run(f"{no_bal}\n")

        # Add payment section for secondary individual if this is the last year
        add_payment_section(doc, federal_owing, quebec_owing if province == "Qu\u00e9bec" else 0, year, last_year)

        if "carryforward_amounts" in secondary_summary and secondary_summary["carryforward_amounts"]:
            carryforward_amounts(doc, secondary_summary, cfg=cfg)


# Section: Crédit de solidarité
def solidarity_creditFR(doc, return_summary, year, last_year, cfg=None):
    para = doc.add_paragraph()

    # Add title with bold and underline
    sol_title = get_text(cfg, 'solidarityTitle', 'Cr\u00e9dit de solidarit\u00e9 :')
    title_run = para.add_run(f'{sol_title}\n')
    apply_style(title_run, get_style(cfg, 'solidarityTitle') or {'bold': True, 'underline': True})

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    # Extract the total solidarity credit amount
    solidarity_amount = return_summary["solidarity_amounts"]["solidarity_credit_amount"]

    # Add the formatted total amount to the paragraph
    para.add_run(f'Vous allez recevoir un total de ')
    total_run = para.add_run(format_currency(solidarity_amount))
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' pour le ')
    para.add_run(f'Cr\u00e9dit de solidarit\u00e9').bold = True

    if year == last_year:
        para.add_run(f' de la fa\u00e7on suivante :\n')
    else:
        para.add_run(f'.\n')

    # Extract amounts for each month
    amounts = {
        "Juillet": return_summary["solidarity_amounts"]["july_amount"],
        "Ao\u00fbt": return_summary["solidarity_amounts"]["august_amount"],
        "Septembre": return_summary["solidarity_amounts"]["september_amount"],
        "Octobre": return_summary["solidarity_amounts"]["october_amount"],
        "Novembre": return_summary["solidarity_amounts"]["november_amount"],
        "D\u00e9cembre": return_summary["solidarity_amounts"]["december_amount"],
        "Janvier": return_summary["solidarity_amounts"]["january_amount"],
        "F\u00e9vrier": return_summary["solidarity_amounts"]["february_amount"],
        "Mars": return_summary["solidarity_amounts"]["march_amount"],
        "Avril": return_summary["solidarity_amounts"]["april_amount"],
        "Mai": return_summary["solidarity_amounts"]["may_amount"],
        "Juin": return_summary["solidarity_amounts"]["june_amount"],
    }

    # Filter out months with zero amounts
    filtered_months = {month: amount for month, amount in amounts.items() if amount > 0}

    if year == last_year:
        if len(filtered_months) <= 4:
            for month, amount in filtered_months.items():
                year = year_plus1 if month in ["Juillet", "Ao\u00fbt", "Septembre", "Octobre", "Novembre", "D\u00e9cembre"] else year_plus2
                para.add_run(f'{month} {year} : {format_currency(amount)}\n')
        else:
            unique_amounts = set(filtered_months.values())

            if len(unique_amounts) == 1:
                amount = next(iter(unique_amounts))
                para.add_run(f'{format_currency(amount)} /mois de Juillet {year_plus1} \u00e0 Juin {year_plus2}\n')
            else:
                amount_ranges = []
                current_amount = None
                start_month = None

                for month, amount in filtered_months.items():
                    year = year_plus1 if month in ["Juillet", "Ao\u00fbt", "Septembre", "Octobre", "Novembre", "D\u00e9cembre"] else year_plus2
                    if current_amount is None:
                        current_amount = amount
                        start_month = month
                    elif abs(current_amount - amount) > 0.3:
                        end_year = year_plus1 if start_month in ["Juillet", "Ao\u00fbt", "Septembre", "Octobre", "Novembre", "D\u00e9cembre"] else year_plus2
                        amount_ranges.append((start_month, month, current_amount, end_year))
                        current_amount = amount
                        start_month = month
                    if month == "Juin":
                        end_year = year_plus2
                        amount_ranges.append((start_month, month, current_amount, end_year))

                for start, end, amount, end_year in amount_ranges:
                    para.add_run(f'{format_currency(amount)} /mois de {start} {year_plus1} \u00e0 {end} {end_year}\n')


# Section: Crédits ACEBE
def ecgeb_creditFR(doc, return_summary, isNewcomer, year, last_year, first_year, cfg=None):
    para = doc.add_paragraph()
    ecgeb_title_text = get_text(cfg, 'ecgebTitle', "Allocation canadienne pour l\u2019\u00e9picerie :")
    title_run = para.add_run(ecgeb_title_text)
    apply_style(title_run, get_style(cfg, 'ecgebTitle') or {'bold': True, 'underline': True})

    if isNewcomer and year == first_year:
        title_run = para.add_run('* ')
        title_run.bold = True
        title_run.underline = True
    title_run = para.add_run('\n')
    title_run.bold = True
    title_run.underline = True

    ecgeb_credit_amount = return_summary["ecgeb_amounts"]["ecgeb_credit_amount"]
    july_amount = return_summary["ecgeb_amounts"].get("july_amount", 0)
    october_amount = return_summary["ecgeb_amounts"].get("october_amount", 0)
    january_amount = return_summary["ecgeb_amounts"].get("january_amount", 0)
    april_amount = return_summary["ecgeb_amounts"].get("april_amount", 0)

    para.add_run(f'Vous allez recevoir un total de ')
    total_run = para.add_run(format_currency(ecgeb_credit_amount))
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' pour la ')
    para.add_run("Allocation canadienne pour l'\u00e9picerie et les besoins essentiels (ACEBE)").bold = True

    if year == last_year:
        para.add_run(f' de la fa\u00e7on suivante :\n')
    else:
        para.add_run(f'.\n')

    if year == last_year:
        if july_amount > 0:
            para.add_run(f'Juillet {year} : {format_currency(july_amount)}\n')
        if october_amount > 0:
            para.add_run(f'Octobre {year} : {format_currency(october_amount)}\n')
        if january_amount > 0:
            para.add_run(f'Janvier {year + 1} : {format_currency(january_amount)}\n')
        if april_amount > 0:
            para.add_run(f'Avril {year + 1} : {format_currency(april_amount)}\n')

    if isNewcomer and year == first_year:
        para.add_run(
            "\n*Notez que vous allez recevoir une lettre de l'Agence du Revenu du Canada vous demandant de fournir vos revenus avant votre arriv\u00e9e au Canada (du 1er janvier jusqu\u2019\u00e0 la date d'arriv\u00e9e). "
            "M\u00eame si cela a \u00e9t\u00e9 mentionn\u00e9 sur la d\u00e9claration, vous devez quand m\u00eame r\u00e9pondre \u00e0 la lettre. Si vous ne r\u00e9pondez pas, ils ne paieront pas le montant de la ACEBE.\n"
        ).italic = True


# Section: Crédit TPS
def gst_creditFR(doc, return_summary, isNewcomer, year, last_year, first_year, cfg=None):
    para = doc.add_paragraph()
    gst_title_text = get_text(cfg, 'gstTitle', 'Cr\u00e9dits TPS/TVH :')
    title_run = para.add_run(gst_title_text)
    apply_style(title_run, get_style(cfg, 'gstTitle') or {'bold': True, 'underline': True})

    if isNewcomer and year == first_year:
        title_run = para.add_run('* ')
        title_run.bold = True
        title_run.underline = True
    title_run = para.add_run('\n')
    title_run.bold = True
    title_run.underline = True

    gst_credit_amount = return_summary["gst_amounts"]["gst_credit_amount"]
    july_amount = return_summary["gst_amounts"].get("july_amount", 0)
    october_amount = return_summary["gst_amounts"].get("october_amount", 0)
    january_amount = return_summary["gst_amounts"].get("january_amount", 0)
    april_amount = return_summary["gst_amounts"].get("april_amount", 0)

    para.add_run(f'Vous allez recevoir un total de ')
    total_run = para.add_run(format_currency(gst_credit_amount))
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' pour les cr\u00e9dits ')
    para.add_run(f'TPS').bold = True

    if year == last_year:
        para.add_run(f' de la fa\u00e7on suivante :\n')
    else:
        para.add_run(f'.\n')

    if year == last_year:
        if july_amount > 0:
            para.add_run(f'Juillet {year} : {format_currency(july_amount)}\n')
        if october_amount > 0:
            para.add_run(f'Octobre {year} : {format_currency(october_amount)}\n')
        if january_amount > 0:
            para.add_run(f'Janvier {year + 1} : {format_currency(january_amount)}\n')
        if april_amount > 0:
            para.add_run(f'Avril {year + 1} : {format_currency(april_amount)}\n')

    if isNewcomer and year == first_year:
        para.add_run(
            "\n*Notez que vous allez recevoir une lettre de l'Agence du Revenu du Canada vous demandant de fournir vos revenus avant votre arriv\u00e9e au Canada (du 1er janvier jusqu\u2019\u00e0 la date d'arriv\u00e9e). "
            "M\u00eame si cela a \u00e9t\u00e9 mentionn\u00e9 sur la d\u00e9claration, vous devez quand m\u00eame r\u00e9pondre \u00e0 la lettre. Si vous ne r\u00e9pondez pas, ils ne paieront pas le montant de la TPS.\n"
        ).italic = True


def carryforward_amountsFR(doc, return_summary, cfg=None):
    para = doc.add_paragraph()
    tuition_title = get_text(cfg, 'tuitionTitle', 'Vos Frais de scolarit\u00e9 report\u00e9s aux ann\u00e9es futures sont :')
    title_run = para.add_run(f'{tuition_title}\n')
    apply_style(title_run, get_style(cfg, 'tuitionTitle') or {'bold': True})

    federal_tuition_amount = return_summary["carryforward_amounts"]["federal_tuition_amount"]
    quebec_tuition_8_percent = return_summary["carryforward_amounts"]["quebec_tuition_8_percent"]

    # Ensure amounts are numeric (float or int)
    if isinstance(federal_tuition_amount, str):
        if federal_tuition_amount.strip():  # Check if not empty
            federal_tuition_amount = int(federal_tuition_amount)
        else:
            federal_tuition_amount = 0

    if isinstance(quebec_tuition_8_percent, str):
        if quebec_tuition_8_percent.strip():  # Check if not empty
            quebec_tuition_8_percent = int(quebec_tuition_8_percent)
        else:
            quebec_tuition_8_percent = 0

    # Format the amounts without decimals and replace commas with spaces
    formatted_federal_amount = locale.format_string("%d", federal_tuition_amount, grouping=True).replace(",", " ")
    formatted_quebec_amount = locale.format_string("%d", quebec_tuition_8_percent, grouping=True).replace(",", " ")

    # Add formatted amounts to the paragraph
    if federal_tuition_amount > 0:
        fed_label = get_text(cfg, 'tuitionFedLabel', 'F\u00e9d\u00e9ral (admissible \u00e0 15%) : ')
        para.add_run(fed_label)
        para.add_run(f"{formatted_federal_amount} $")
    if quebec_tuition_8_percent > 0:
        qc_label = get_text(cfg, 'tuitionQCLabel', 'QC (admissible \u00e0 8%) : ')
        para.add_run(f'\n{qc_label}')
        para.add_run(f"{formatted_quebec_amount} $")

    # Add additional explanatory text
    tuition_exp = get_text(cfg, 'tuitionExplanation',
        'Ces frais de scolarit\u00e9 accumul\u00e9s sont des cr\u00e9dits d\'imp\u00f4t que vous allez utiliser dans vos futures d\u00e9clarations, lorsque vous g\u00e9n\u00e9rez un revenu et payez des imp\u00f4ts.')
    r = para.add_run(f'\n\n{tuition_exp}')
    apply_style(r, get_style(cfg, 'tuitionExplanation') or {'italic': True})




# Section: Prestations pour Enfants
def child_benefitFR(doc, return_summary, year, last_year, cfg=None):
    para = doc.add_paragraph()
    # Add title directly with bold and underline
    ccb_title = get_text(cfg, 'ccbTitle', 'Allocation canadienne pour enfants (ACE) :')
    title_run = para.add_run(f'{ccb_title}\n')
    apply_style(title_run, get_style(cfg, 'ccbTitle') or {'bold': True, 'underline': True})

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    # Extract total Canada Child Benefit amount
    total_ccb = return_summary["ccb_amounts"]["ccb_amount"]
    amounts = {
        f"Juillet {year_plus1}": return_summary["ccb_amounts"]["july_amount"],
        f"Ao\u00fbt {year_plus1}": return_summary["ccb_amounts"]["august_amount"],
        f"Septembre {year_plus1}": return_summary["ccb_amounts"]["september_amount"],
        f"Octobre {year_plus1}": return_summary["ccb_amounts"]["october_amount"],
        f"Novembre {year_plus1}": return_summary["ccb_amounts"]["november_amount"],
        f"D\u00e9cembre {year_plus1}": return_summary["ccb_amounts"]["december_amount"],
        f"Janvier {year_plus2}": return_summary["ccb_amounts"]["january_amount"],
        f"F\u00e9vrier {year_plus2}": return_summary["ccb_amounts"]["february_amount"],
        f"Mars {year_plus2}": return_summary["ccb_amounts"]["march_amount"],
        f"Avril {year_plus2}": return_summary["ccb_amounts"]["april_amount"],
        f"Mai {year_plus2}": return_summary["ccb_amounts"]["may_amount"],
        f"Juin {year_plus2}": return_summary["ccb_amounts"]["june_amount"],
    }

    # Display total benefit amount
    para.add_run(f'Vous allez recevoir un total de ').bold = False
    total_run = para.add_run(format_currency(total_ccb))
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' pour ')
    para.add_run(f'l\u2019Allocation canadienne pour enfants ').bold = True
    if year == last_year:
        para.add_run(f' de la fa\u00e7on suivante :\n')
    else:
        para.add_run(f'.\n')

    # Filter out months with zero amounts
    filtered_months = {month: amount for month, amount in amounts.items() if amount > 0}

    if year == last_year:
        if len(filtered_months) <= 4:
            for month, amount in filtered_months.items():
                para.add_run(f'{month} : {format_currency(amount)}\n')
        else:
            # Check for consistency of amounts for the range
            unique_amounts = set(filtered_months.values())

            if len(unique_amounts) == 1:
                # If all amounts are the same
                amount = next(iter(unique_amounts))  # Get the single value
                para.add_run(f'{format_currency(amount)} /mois de Juillet {year_plus1} \u00e0 Juin {year_plus2}')
            else:
                # Group amounts by ranges
                amount_ranges = []
                current_amount = None
                start_month = None

                for month, amount in filtered_months.items():
                    # Automatically determine the year from the month
                    year = month.split()[1]  # Extract year from month string

                    if current_amount is None:
                        current_amount = amount
                        start_month = month
                    elif abs(current_amount - amount) > 0.3:
                        amount_ranges.append((start_month, month, current_amount, year))
                        current_amount = amount
                        start_month = month

                # Handle the last range
                if start_month:
                    amount_ranges.append((start_month, month, current_amount, year))

                for start, end, amount, year in amount_ranges:
                    para.add_run(f'{format_currency(amount)} /mois de {start} \u00e0 {end}')


# Section: Allocation Famille (Québec)
def quebec_family_allowanceFR(doc, return_summary, year, last_year, cfg=None):
    para = doc.add_paragraph()

    fa_title = get_text(cfg, 'familyAllowanceTitle', 'Allocation Famille :')
    title_run = para.add_run(fa_title + '\n')
    apply_style(title_run, get_style(cfg, 'familyAllowanceTitle') or {'bold': True, 'underline': True})

    total_allowance_amount = return_summary["family_allowance_amounts"]["fa_amount"]

    para.add_run(f'Vous allez recevoir un total de ').bold = False
    total_run = para.add_run(format_currency(total_allowance_amount))
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' pour ')
    fa_title = get_text(cfg, 'familyAllowanceTitle', 'Allocation Famille :')
    para.add_run(fa_title).bold = True
    if year == last_year:
        para.add_run(f' de la fa\u00e7on suivante :\n')
    else:
        para.add_run(f'.\n')

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    july_amount = return_summary["family_allowance_amounts"]["july_amount"]
    october_amount = return_summary["family_allowance_amounts"]["october_amount"]
    january_amount = return_summary["family_allowance_amounts"]["january_amount"]
    april_amount = return_summary["family_allowance_amounts"]["april_amount"]

    if year == last_year:
        if july_amount > 0:
            para.add_run(f'Juillet {year_plus1} : {format_currency(july_amount)}\n')
        if october_amount > 0:
            para.add_run(f'Octobre {year_plus1} : {format_currency(october_amount)}\n')
        if january_amount > 0:
            para.add_run(f'Janvier {year_plus2} : {format_currency(january_amount)}\n')
        if april_amount > 0:
            para.add_run(f'Avril {year_plus2} : {format_currency(april_amount)}\n')


def conclusionFR(doc, isCouple=False, cfg=None):
    para = doc.add_paragraph()

    # Add the red text above the conclusion
    if isCouple:
        waiting_default = 'Nous attendons les formulaires d\u2019autorisations sign\u00e9s pour soumettre vos d\u00e9clarations.'
    else:
        waiting_default = 'Nous attendons les formulaire d\u2019autorisations sign\u00e9 pour soumettre votre d\u00e9claration.'
    waiting_text = get_text(cfg, 'conclusionWaiting', waiting_default)
    red_run = para.add_run(f'\n{waiting_text}\n')
    apply_style(red_run, get_style(cfg, 'conclusionWaiting') or {'color': '#cd3350'})

    # Add normal text for "Thank you."
    thank_text = get_text(cfg, 'thankYou', 'Merci')
    para.add_run(f'\n{thank_text}\n')

    # Add the main conclusion text with specified formatting
    disclaimer_default = (
        "Nous, \u00e0 Sankari Inc., sommes heureux de r\u00e9pondre \u00e0 vos demandes de renseignements fiscaux et / ou de produire vos d\u00e9clarations de revenus en "
        "fonction des renseignements que vous nous fournissez. Les informations inexactes ou incompl\u00e8tes fournies par vous peuvent conduire \u00e0 des "
        "conseils inad\u00e9quats ou incorrects pour lesquels l\u2019\u00e9quipe de Sankari Inc. ne peut \u00eatre tenue pour responsable. Vous, le client est responsable de "
        "fournir l'information et la documentation correcte \u00e0 l'\u00e9quipe de Sankari Inc."
    )
    disclaimer_text = get_text(cfg, 'disclaimer', disclaimer_default)
    conclusion_run = para.add_run(f'\n{disclaimer_text}')
    apply_style(conclusion_run, get_style(cfg, 'disclaimer') or {'fontSize': 8, 'color': '#7f7f7f', 'italic': True})
