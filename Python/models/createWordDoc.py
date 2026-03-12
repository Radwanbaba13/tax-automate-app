from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import locale
import time

from models.docTextHelper import resolve_doc_type_key, get_cfg, get_text, get_style, apply_style, apply_alignment, styled_run


def createIndividualWordDoc(individual, output_file_path, doc_text_config=None):
    if individual['language'] == 'EN':
        createIndividualWordDocEN(individual, output_file_path, doc_text_config)
    if individual['language'] == 'FR':
        createIndividualWordDocFR(individual, output_file_path, doc_text_config)

def createCoupleWordDoc(couple_summaries, output_file_path, doc_text_config=None):
    primary_client = next((c for c in couple_summaries if c['isPrimary']), None)
    language = primary_client['language']
    if language == 'EN':
        createCoupleWordDocEN(couple_summaries, output_file_path, doc_text_config)
    elif language == 'FR':
        createCoupleWordDocFR(couple_summaries, output_file_path, doc_text_config)


try:
    locale.setlocale(locale.LC_ALL, 'fr_CA.UTF-8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_ALL, 'French_Canada.1252')  # Windows
    except locale.Error:
        pass

def format_currency(amount):
    # Convert the amount to a float and format it as a string
    return f"{locale.format_string('%.2f', amount, grouping=True).replace('.', ',')} $"

# Function to set default font and size for the entire document
def set_default_font(doc, font_name, font_size):
    styles = doc.styles['Normal']
    font = styles.font
    font.name = font_name
    font.size = Pt(font_size)
    # Remove default paragraph spacing so sections don't have large gaps
    styles.paragraph_format.space_before = Pt(0)
    styles.paragraph_format.space_after = Pt(0)

# Function that adds a hyperlink to a paragraph.
def add_hyperlink(paragraph, text, url):
    # Create the w:hyperlink tag and add required attributes
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    italic = OxmlElement('w:i')
    rPr.append(italic)
    # Apply properties to the run
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

# Function to create a new document and set global styles
def createIndividualWordDocEN(individual, output_file_path, doc_text_config=None):
    return_summary = individual['summary']
    year = individual['year']
    ind_title = individual['title']
    isMailQC = individual['isMailQC']
    isNewcomer = individual['isNewcomer']
    doc = Document()

    cfg = get_cfg(doc_text_config, resolve_doc_type_key('EN', is_couple=False))

    # Set default font to Calibri and size 10pt for the entire document
    set_default_font(doc, "Calibri", 10)

    # Check for each section and call the respective function only if the section exists
    section_1(doc, return_summary, ind_title, year, isMailQC, cfg=cfg)
    tax_return(doc, return_summary, year, cfg=cfg)

    if "gst_amounts" in return_summary and return_summary["gst_amounts"]:
        gst_credit(doc, return_summary, isNewcomer, year, cfg=cfg)

    if "ecgeb_amounts" in return_summary and return_summary["ecgeb_amounts"]:
        ecgeb_credit(doc, return_summary, isNewcomer, year, cfg=cfg)

    if "solidarity_amounts" in return_summary and return_summary["solidarity_amounts"]:
        solidarity_credit(doc, return_summary, year, cfg=cfg)

    if "carbon_rebate_amounts" in return_summary and return_summary["carbon_rebate_amounts"]:
        carbon_rebate(doc, return_summary,  year, cfg=cfg)

    if "climate_action_credit_amounts" in return_summary and return_summary["climate_action_credit_amounts"]:
        climate_action_credit(doc, return_summary, year, cfg=cfg)

    if "ontario_trillium_amounts" in return_summary and return_summary["ontario_trillium_amounts"]:
        ontario_trillium_benefit(doc, return_summary["ontario_trillium_amounts"], year, cfg=cfg)

    if "ccb_amounts" in return_summary and return_summary["ccb_amounts"]:
        child_benefit(doc, return_summary, year, cfg=cfg)

    if "family_allowance_amounts" in return_summary and return_summary["family_allowance_amounts"]:
        quebec_family_allowance(doc, return_summary, year, cfg=cfg)

    if "carryforward_amounts" in return_summary and return_summary["carryforward_amounts"]:
        carryforward_amounts(doc, return_summary, cfg=cfg)


    conclusion(doc, isMailQC, cfg=cfg)

    # Save the document in the specified path
    doc.save(output_file_path)

def createCoupleWordDocEN(couple_summaries, output_file_path, doc_text_config=None):
    doc = Document()

    cfg = get_cfg(doc_text_config, resolve_doc_type_key('EN', is_couple=True))

    # Set default font to Calibri and size 10pt for the entire document
    set_default_font(doc, "Calibri", 10)

    # Initialize variables to track the year, isMailQC, and isNewcomer
    year = couple_summaries[0]['year']
    isMailQC = False
    isNewcomer = False

    for individual in couple_summaries:
        if individual['isMailQC']:
            isMailQC = True
        if individual['isNewcomer']:
            isNewcomer = True

    # Process the primary individual
    primary_individual = next(individual for individual in couple_summaries if individual['isPrimary'])
    secondary_individual = next(individual for individual in couple_summaries if not individual['isPrimary'])

    section_1(doc, primary_individual['summary'], primary_individual['title'], year, isMailQC, couple=True, secondary_summary=secondary_individual['summary'], secondary_ind_title=secondary_individual['title'], cfg=cfg)
    tax_return(doc, primary_individual['summary'], year,  secondary_individual['summary'], primary_individual['title'], secondary_individual['title'], isCouple=True, cfg=cfg)

    for individual in couple_summaries:
        return_summary = individual['summary']
        if "gst_amounts" in return_summary and return_summary["gst_amounts"]:
            gst_credit(doc, return_summary, isNewcomer, year, cfg=cfg)

        if "ecgeb_amounts" in return_summary and return_summary["ecgeb_amounts"]:
            ecgeb_credit(doc, return_summary, isNewcomer, year, cfg=cfg)

        if "carbon_rebate_amounts" in return_summary and return_summary["carbon_rebate_amounts"]:
            carbon_rebate(doc, return_summary,  year, cfg=cfg)

        if "ontario_trillium_amounts" in return_summary and return_summary["ontario_trillium_amounts"]:
            ontario_trillium_benefit(doc, return_summary["ontario_trillium_amounts"], year, cfg=cfg)

        if "climate_action_credit_amounts" in return_summary and return_summary["climate_action_credit_amounts"]:
            climate_action_credit(doc, return_summary, year, cfg=cfg)

        if "solidarity_amounts" in return_summary and return_summary["solidarity_amounts"]:
            solidarity_credit(doc, return_summary, year, cfg=cfg)

        if "ccb_amounts" in return_summary and return_summary["ccb_amounts"]:
            child_benefit(doc, return_summary, year, cfg=cfg)

        if "family_allowance_amounts" in return_summary and return_summary["family_allowance_amounts"]:
            quebec_family_allowance(doc, return_summary, year, cfg=cfg)


    conclusion(doc, cfg=cfg)

    # Save the document in the specified path
    doc.save(output_file_path)

# Section 1: Title and Header
def section_1(doc, primary_summary, ind_title, year, isMailQC, couple=False, secondary_summary=None, secondary_ind_title=None, cfg=None):
   # Title: Centered, Dark Gray, 14pt, Calibri
    default_title = f"Summary of your {year} Tax {'Declarations' if couple else 'Declaration'}"
    title_text = get_text(cfg, 'docTitle', default_title, year=year)
    title = doc.add_paragraph(title_text)
    title_style = get_style(cfg, 'docTitle') or {'fontSize': 14, 'color': '#414141'}
    apply_alignment(title, title_style)
    if not get_style(cfg, 'docTitle'):
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    apply_style(run, title_style)

    # Recipient Name for Primary Individual
    para = doc.add_paragraph()
    primary_first_name = primary_summary["tax_summary"]["first_name"]
    primary_last_name = primary_summary["tax_summary"]["last_name"]
    province = primary_summary["tax_summary"]["province"]
    para.add_run(f"{ind_title}. {primary_last_name}")

    # Display Secondary Individual’s name if applicable
    if couple and secondary_summary and secondary_ind_title:
        secondary_first_name = secondary_summary["tax_summary"]["first_name"]
        secondary_last_name = secondary_summary["tax_summary"]["last_name"]
        para.add_run(f" & {secondary_ind_title}. {secondary_last_name}")

    # Introduction paragraph
    para = doc.add_paragraph()
    default_attach = f'We have attached all the documents related to your {year} tax {"declarations" if couple else "declaration"}.'
    styled_run(para, cfg, 'introAttachment', default_attach + '\n', {'bold': False}, year=year)
    styled_run(para, cfg, 'introPassword', 'The password consists of the nine digits of your Social Insurance Number.\n', {'bold': True})
    styled_run(para, cfg, 'introCopyDescription', 'The document named COPY is a copy of your complete tax return. ', {'bold': False})
    styled_run(para, cfg, 'introCopyNoPrint', 'You do not need to print it or sign it; ', {'bold': True})
    styled_run(para, cfg, 'introCopyKeep', 'please keep it for your records and review it carefully to ensure everything is accurate and complete.', {'bold': False})

    # "Very Important" without extra paragraph spacing
    vi_text = get_text(cfg, 'veryImportantHeading', '** Very Important:')
    very_important_run = para.add_run(f'\n\n{vi_text}\n')
    apply_style(very_important_run, get_style(cfg, 'veryImportantHeading') or {'bold': True, 'underline': True, 'color': '#cd3350'})

    # Only include Québec-specific instructions if the province is Québec
    if province == "Quebec":
        if isMailQC:
            # Special message if isMailQC is True
            fed_title = get_text(cfg, 'qcMailFedTitle', 'Regarding your Federal tax return:')
            federal_title_run = para.add_run(f'{fed_title}\n')
            apply_style(federal_title_run, get_style(cfg, 'qcMailFedTitle') or {'bold': True, 'underline': True})

            fed_not_submitted = get_text(cfg, 'qcMailFedNotSubmitted', 'Please be advised that your Federal tax return has not been submitted to the government yet.')
            r = para.add_run(f'{fed_not_submitted}\n')
            apply_style(r, get_style(cfg, 'qcMailFedNotSubmitted') or {'bold': True})

            name = f"{primary_first_name} {primary_last_name}"
            sec_name = f"{secondary_first_name} {secondary_last_name}" if couple and secondary_summary else ""
            auth_text = get_text(cfg, 'qcMailFedAuthForm',
                f'Attached {"are two Authorization Forms" if couple else "is an Authorization Form"}. Please e-sign it (or print & sign) and e-mail it back to us as soon as possible so we can EFILE your return.')
            para.add_run(f'{auth_text}\n')

            sign_text = get_text(cfg, 'qcMailSignPartF', 'Please sign Part F.')
            para.add_run(f'{sign_text}\n\n')

            # Québec tax return
            qc_title = get_text(cfg, 'qcMailQCTitle', 'Regarding your Québec tax return:')
            quebec_title_run = para.add_run(f'{qc_title}\n')
            apply_style(quebec_title_run, get_style(cfg, 'qcMailQCTitle') or {'bold': True, 'underline': True})

            cannot_efile = get_text(cfg, 'qcMailQCCannotEfile', 'Please note that your Québec tax return cannot be transmitted via Efile.')
            para.add_run(f'{cannot_efile}\n')

            print_text = get_text(cfg, 'qcMailQCPrint',
                f'For that reason, you need to print the document "QC {year} - {name}.pdf", sign at the bottom of page ##, and mail it to the following address:',
                year=year, name=name, primaryName=name, secondaryName=sec_name)
            para.add_run(f'{print_text}\n')

            # Address in italics
            addr_text = get_text(cfg, 'qcAddress', 'Revenu Québec\nC. P. 2500, succursale Place-Desjardins\nMontréal (Québec) H5B 1A3')
            address = para.add_run(f'{addr_text}\n\n')
            apply_style(address, get_style(cfg, 'qcAddress') or {'italic': True})

            on_behalf = get_text(cfg, 'qcMailOnBehalf',
                '*If you would like us to mail the declaration on your behalf, please email us the signed declaration (e-signature that looks like your signature), and we will print and mail it by registered mail (with a tracking number) to QC Revenue. Please note that there will be an additional service fee of $25 plus Canada Post fees.')
            para.add_run(f'{on_behalf}\n')
        else:
            # Original message if isMailQC is False
            not_submitted = get_text(cfg, 'qcEfileNotSubmitted', f'Please be advised that your tax {"returns have" if couple else "return has"} not been submitted to the government yet.')
            r = para.add_run(f'{not_submitted}\n')
            apply_style(r, get_style(cfg, 'qcEfileNotSubmitted') or {'bold': True})

            auth_text = get_text(cfg, 'qcEfileAuthForms',
                f'Attached are {"four" if couple else "two"} Authorization Forms. Please e-sign them (or print & sign) and e-mail them back to us as soon as possible so we can EFILE your returns.')
            para.add_run(f'{auth_text}\n')

            sign_fed = get_text(cfg, 'qcEfileSignFed', 'For the Federal Form, please sign Part F.')
            r = para.add_run(f'{sign_fed}\n')
            apply_style(r, get_style(cfg, 'qcEfileSignFed') or {'italic': True})

            sign_qc = get_text(cfg, 'qcEfileSignQC', 'For the Quebec Form, please sign at the end of section 4.')
            r = para.add_run(sign_qc)
            apply_style(r, get_style(cfg, 'qcEfileSignQC') or {'italic': True})
    else:
        # Non-Québec clients (no Québec-related details)
        not_submitted = get_text(cfg, 'nonQcNotSubmitted', f'Please be advised that your tax {"returns have" if couple else "return has"} not been submitted to the government yet.')
        r = para.add_run(f'{not_submitted}\n')
        apply_style(r, get_style(cfg, 'nonQcNotSubmitted') or {'bold': True})

        auth_text = get_text(cfg, 'nonQcAuthForm',
            f'Attached {"are two Authorization Forms" if couple else "is an Authorization Form"}. Please e-sign it (or print & sign) and e-mail it back to us as soon as possible so we can EFILE your return.')
        para.add_run(f'{auth_text}\n')

        sign_text = get_text(cfg, 'nonQcSignPartF', 'Please sign Part F.')
        r = para.add_run(f'{sign_text}\n')
        apply_style(r, get_style(cfg, 'nonQcSignPartF') or {'italic': True})


def tax_return(doc, return_summary, year, secondary_summary=None, primary_title=None, secondary_title=None, isCouple=False, cfg=None):
    province = return_summary["tax_summary"]["province"]

    # Create the section for tax return results
    para = doc.add_paragraph()
    results_text = get_text(cfg, 'resultsHeading', 'RESULTS')
    bold_red_results = para.add_run(f'{results_text}\n')
    apply_style(bold_red_results, get_style(cfg, 'resultsHeading') or {'bold': True, 'color': '#cd3350'})

    # Process primary individual’s tax details
    if isCouple:
        para.add_run(f"\n{primary_title}. {return_summary['tax_summary']['last_name']}\n")

    # Federal Tax Return for primary individual
    fed_label = get_text(cfg, 'federalReturnLabel', 'Federal Tax return')
    r = para.add_run(f'{fed_label}\n')
    apply_style(r, get_style(cfg, 'federalReturnLabel') or {'bold': True})
    federal_refund = return_summary["tax_summary"]["federal_refund"]
    federal_owing = return_summary["tax_summary"]["federal_owing"]

    if federal_refund > 2:
        refund_prefix = get_text(cfg, 'refundPrefix', 'You are entitled to a refund of')
        para.add_run(f"{refund_prefix} ").bold = False
        refund_run = para.add_run(f"${federal_refund:,.2f} \n")
        refund_run.bold = True
        refund_run.font.color.rgb = RGBColor(0, 128, 0)
    elif federal_owing > 2:
        owing_prefix = get_text(cfg, 'owingPrefix', 'You owe the amount of')
        para.add_run(f"{owing_prefix} ").bold = False
        owing_run = para.add_run(f"${federal_owing:,.2f} \n")
        owing_run.bold = True
        owing_run.font.color.rgb = RGBColor(255, 0, 0)
    else:
        no_bal = get_text(cfg, 'noBalance', 'You have no Refund or Balance due.')
        para.add_run(f"{no_bal}\n")

    # Quebec Tax Return for primary individual (only if province is Quebec)
    if province == "Quebec":
        qc_label = get_text(cfg, 'quebecReturnLabel', 'Quebec Tax return')
        r = para.add_run(f'{qc_label}\n')
        apply_style(r, get_style(cfg, 'quebecReturnLabel') or {'bold': True})
        quebec_refund = return_summary["tax_summary"]["quebec_refund"]
        quebec_owing = return_summary["tax_summary"]["quebec_owing"]

        if quebec_refund > 2:
            refund_prefix = get_text(cfg, 'refundPrefix', 'You are entitled to a refund of')
            para.add_run(f"{refund_prefix} ").bold = False
            refund_run = para.add_run(f"${quebec_refund:,.2f}\n")
            refund_run.bold = True
            refund_run.font.color.rgb = RGBColor(0, 128, 0)
        elif quebec_owing > 2:
            owing_prefix = get_text(cfg, 'owingPrefix', 'You owe the amount of')
            para.add_run(f"{owing_prefix} ").bold = False
            owing_run = para.add_run(f"${quebec_owing:,.2f}\n")
            owing_run.bold = True
            owing_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            no_bal = get_text(cfg, 'noBalance', 'You have no Refund or Balance due.')
            para.add_run(f"{no_bal}\n")

    # Add payment section for primary individual
    quebec_owing = locals().get('quebec_owing', 0)

    if federal_owing > 2 or quebec_owing > 2:
        para = doc.add_paragraph()
        if federal_owing > 2 and quebec_owing > 2:
            owing_text = get_text(cfg, 'paymentOwingFedAndQC', 'You owe an amount on your Federal and Quebec returns;')
        elif federal_owing > 2:
            owing_text = get_text(cfg, 'paymentOwingFed', 'You owe an amount on your Federal return;')
        else:
            owing_text = get_text(cfg, 'paymentOwingQC', 'You owe an amount on your Quebec return;')
        r = para.add_run(f'{owing_text} ')
        apply_style(r, get_style(cfg, 'paymentOwingFed') or {'italic': True})

        deadline = get_text(cfg, 'paymentDeadline', f'please make sure to pay the balance due by April 30, {int(year) + 1}, to avoid paying any interest.', dueYear=str(int(year) + 1))
        r = para.add_run(f'{deadline} ')
        apply_style(r, get_style(cfg, 'paymentDeadline') or {'italic': True})

        howto = get_text(cfg, 'paymentHowTo', 'Please wait a few days after we E-file to pay your outstanding balance. For more details on how to pay the amount due, please click on:')
        r = para.add_run(f'{howto} ')
        apply_style(r, get_style(cfg, 'paymentHowTo') or {'italic': True})

        # Add links for payment based on owing status
        if federal_owing > 2:
            add_hyperlink(para, 'Federal', 'https://www.canada.ca/en/revenue-agency/services/payments/payments-cra.html')

        if federal_owing > 2 and quebec_owing > 2:
            para.add_run(' and ')
        if quebec_owing > 2:
            add_hyperlink(para, 'Quebec', 'https://www.revenuquebec.ca/en/citizens/income-tax-return/paying-a-balance-due-or-receiving-a-refund/income-tax-balance-due/')
        para.add_run('\n')

    if isCouple and "carryforward_amounts" in return_summary and return_summary["carryforward_amounts"]:
            para = doc.add_paragraph()
            tuition_title = get_text(cfg, 'tuitionTitle', 'Your accumulated tuition fees carried forward to future years:')
            title_run = para.add_run(f'{tuition_title}\n')
            apply_style(title_run, get_style(cfg, 'tuitionTitle') or {'bold': True})

            federal_tuition_amount = return_summary["carryforward_amounts"]["federal_tuition_amount"]
            quebec_tuition_8_percent = return_summary["carryforward_amounts"]["quebec_tuition_8_percent"]

            # Convert string values safely (removing commas for conversion)
            if isinstance(federal_tuition_amount, str):
                if federal_tuition_amount.strip():
                    federal_tuition_amount = int(federal_tuition_amount.replace(",", ""))
                else:
                    federal_tuition_amount = 0
            if isinstance(quebec_tuition_8_percent, str):
                if quebec_tuition_8_percent.strip():
                    quebec_tuition_8_percent = int(quebec_tuition_8_percent.replace(",", ""))
                else:
                    quebec_tuition_8_percent = 0

            # Format numbers with commas as thousand separators when adding them to the document
            if federal_tuition_amount > 0:
                fed_label = get_text(cfg, 'tuitionFedLabel', 'Federal (eligible to 15%):  $')
                para.add_run(fed_label)
                para.add_run(f"{federal_tuition_amount:,}")
            if 'quebec_tuition_8_percent' in locals() and quebec_tuition_8_percent is not None and quebec_tuition_8_percent > 0:
                qc_label = get_text(cfg, 'tuitionQCLabel', 'QC (eligible to 8%): $')
                para.add_run(f'\n{qc_label}')
                para.add_run(f"{quebec_tuition_8_percent:,}")

            tuition_exp = get_text(cfg, 'tuitionExplanation', 'Those accumulated tuition fees are tax credits that you will be using in future tax declarations when you work and pay tax on your income.')
            r = para.add_run(f'\n\n{tuition_exp}')
            apply_style(r, get_style(cfg, 'tuitionExplanation') or {'italic': True})

    # Add secondary individual if isCouple is True and secondary_summary is provided
    if isCouple and secondary_summary:
        para.add_run(f"\n{secondary_title}. {secondary_summary['tax_summary']['last_name']}\n")

        # Federal Tax Return for secondary individual
        fed_label = get_text(cfg, 'federalReturnLabel', 'Federal Tax return')
        r = para.add_run(f'{fed_label}\n')
        apply_style(r, get_style(cfg, 'federalReturnLabel') or {'bold': True})
        federal_refund = secondary_summary["tax_summary"]["federal_refund"]
        federal_owing = secondary_summary["tax_summary"]["federal_owing"]

        if federal_refund > 2:
            refund_prefix = get_text(cfg, 'refundPrefix', 'You are entitled to a refund of')
            para.add_run(f"{refund_prefix} ").bold = False
            refund_run = para.add_run(f"${federal_refund:,.2f}\n")
            refund_run.bold = True
            refund_run.font.color.rgb = RGBColor(0, 128, 0)
        elif federal_owing > 2:
            owing_prefix = get_text(cfg, 'owingPrefix', 'You owe the amount of')
            para.add_run(f"{owing_prefix} ").bold = False
            owing_run = para.add_run(f"${federal_owing:,.2f}\n")
            owing_run.bold = True
            owing_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            no_bal = get_text(cfg, 'noBalance', 'You have no Refund or Balance due.')
            para.add_run(f"{no_bal}\n")

        # Quebec Tax Return for secondary individual (only if province is Quebec)
        if province == "Quebec":
            qc_label = get_text(cfg, 'quebecReturnLabel', 'Quebec Tax return')
            r = para.add_run(f'{qc_label}\n')
            apply_style(r, get_style(cfg, 'quebecReturnLabel') or {'bold': True})
            quebec_refund = secondary_summary["tax_summary"]["quebec_refund"]
            quebec_owing = secondary_summary["tax_summary"]["quebec_owing"]

            if quebec_refund > 2:
                refund_prefix = get_text(cfg, 'refundPrefix', 'You are entitled to a refund of')
                para.add_run(f"{refund_prefix} ").bold = False
                refund_run = para.add_run(f"${quebec_refund:,.2f}")
                refund_run.bold = True
                refund_run.font.color.rgb = RGBColor(0, 128, 0)
            elif quebec_owing > 2:
                owing_prefix = get_text(cfg, 'owingPrefix', 'You owe the amount of')
                para.add_run(f"{owing_prefix} ").bold = False
                owing_run = para.add_run(f"${quebec_owing:,.2f}")
                owing_run.bold = True
                owing_run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                no_bal = get_text(cfg, 'noBalance', 'You have no Refund or Balance due.')
                para.add_run(f"{no_bal}\n")

        # Add payment section for secondary individual
        quebec_owing = locals().get('quebec_owing', 0)
        if federal_owing > 2 or quebec_owing > 2:

            para = doc.add_paragraph()
            if federal_owing > 2 and quebec_owing > 2:
                owing_text = get_text(cfg, 'paymentOwingFedAndQC', 'You owe an amount on your Federal and Quebec returns;')
            elif federal_owing > 2:
                owing_text = get_text(cfg, 'paymentOwingFed', 'You owe an amount on your Federal return;')
            else:
                owing_text = get_text(cfg, 'paymentOwingQC', 'You owe an amount on your Quebec return;')
            r = para.add_run(f'{owing_text} ')
            apply_style(r, get_style(cfg, 'paymentOwingFed') or {'italic': True})

            deadline = get_text(cfg, 'paymentDeadline', f'please make sure to pay the balance due by April 30, {int(year) + 1}, to avoid paying any interest.', dueYear=str(int(year) + 1))
            r = para.add_run(f'{deadline} ')
            apply_style(r, get_style(cfg, 'paymentDeadline') or {'italic': True})

            howto = get_text(cfg, 'paymentHowTo', 'Please wait a few days after we E-file to pay your outstanding balance. For more details on how to pay the amount due, please click on:')
            r = para.add_run(f'{howto} ')
            apply_style(r, get_style(cfg, 'paymentHowTo') or {'italic': True})

            # Add links for payment based on owing status
            if federal_owing > 2:
                add_hyperlink(para, 'Federal', 'https://www.canada.ca/en/revenue-agency/services/payments/payments-cra.html')

            if federal_owing > 2 and quebec_owing > 2:
                para.add_run(' and ')
            if quebec_owing > 2:
                add_hyperlink(para, 'Quebec', 'https://www.revenuquebec.ca/en/citizens/income-tax-return/paying-a-balance-due-or-receiving-a-refund/income-tax-balance-due/')
            para.add_run('\n')

        if "carryforward_amounts" in secondary_summary and secondary_summary["carryforward_amounts"]:
            para = doc.add_paragraph()
            tuition_title = get_text(cfg, 'tuitionTitle', 'Your accumulated tuition fees carried forward to future years:')
            title_run = para.add_run(f'{tuition_title}\n')
            apply_style(title_run, get_style(cfg, 'tuitionTitle') or {'bold': True})

            federal_tuition_amount = secondary_summary["carryforward_amounts"]["federal_tuition_amount"]
            quebec_tuition_8_percent = secondary_summary["carryforward_amounts"]["quebec_tuition_8_percent"]

            # Convert string values safely (removing commas for conversion)
            if isinstance(federal_tuition_amount, str):
                if federal_tuition_amount.strip():
                    federal_tuition_amount = int(federal_tuition_amount.replace(",", ""))
                else:
                    federal_tuition_amount = 0
            if isinstance(quebec_tuition_8_percent, str):
                if quebec_tuition_8_percent.strip():
                    quebec_tuition_8_percent = int(quebec_tuition_8_percent.replace(",", ""))
                else:
                    quebec_tuition_8_percent = 0

            # Format numbers with commas as thousand separators when adding them to the document
            if federal_tuition_amount > 0:
                fed_label = get_text(cfg, 'tuitionFedLabel', 'Federal (eligible to 15%):  $')
                para.add_run(fed_label)
                para.add_run(f"{federal_tuition_amount:,}")
            if 'quebec_tuition_8_percent' in locals() and quebec_tuition_8_percent is not None and quebec_tuition_8_percent > 0:
                qc_label = get_text(cfg, 'tuitionQCLabel', 'QC (eligible to 8%): $')
                para.add_run(f'\n{qc_label}')
                para.add_run(f"{quebec_tuition_8_percent:,}")

            tuition_exp = get_text(cfg, 'tuitionExplanation', 'Those accumulated tuition fees are tax credits that you will be using in future tax declarations when you work and pay tax on your income.')
            r = para.add_run(f'\n\n{tuition_exp}')
            apply_style(r, get_style(cfg, 'tuitionExplanation') or {'italic': True})


# Section: Solidarity Credit
def solidarity_credit(doc, return_summary, year, cfg=None):
    para = doc.add_paragraph()
    # Add title directly with bold and underline
    sol_title = get_text(cfg, 'solidarityTitle', 'Solidarity Credits:')
    title_run = para.add_run(f'{sol_title}\n')
    apply_style(title_run, get_style(cfg, 'solidarityTitle') or {'bold': True, 'underline': True})

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    solidarity_amount = return_summary["solidarity_amounts"]["solidarity_credit_amount"]

    para.add_run(f'You will receive a total of ').bold = False
    total_run = para.add_run(f"${solidarity_amount:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' for the ')
    para.add_run(f'Solidarity tax credit ').bold = True
    para.add_run(f'as follows:\n')

    # Extract amounts for each month
    amounts = {
        "July": return_summary["solidarity_amounts"]["july_amount"],
        "August": return_summary["solidarity_amounts"]["august_amount"],
        "September": return_summary["solidarity_amounts"]["september_amount"],
        "October": return_summary["solidarity_amounts"]["october_amount"],
        "November": return_summary["solidarity_amounts"]["november_amount"],
        "December": return_summary["solidarity_amounts"]["december_amount"],
        "January": return_summary["solidarity_amounts"]["january_amount"],
        "February": return_summary["solidarity_amounts"]["february_amount"],
        "March": return_summary["solidarity_amounts"]["march_amount"],
        "April": return_summary["solidarity_amounts"]["april_amount"],
        "May": return_summary["solidarity_amounts"]["may_amount"],
        "June": return_summary["solidarity_amounts"]["june_amount"],
    }

    # Filter out months with a zero amount
    filtered_months = {month: amount for month, amount in amounts.items() if amount > 0}

    if len(filtered_months) <= 4:
        for month, amount in filtered_months.items():
            year = year_plus1 if month in ["July", "August", "September", "October", "November", "December"] else year_plus2
            para.add_run(f'{month} {year}: ${amount:,.2f}\n')
    else:
        # Check for consistency of amounts for the range
        unique_amounts = set(filtered_months.values())

        if len(unique_amounts) == 1:
            # If all amounts are the same
            amount = next(iter(unique_amounts))  # Get the single value
            para.add_run(f'${amount:,.2f}/month from July {year_plus1} to June {year_plus2}\n')
        else:
            # Group amounts by ranges
            amount_ranges = []
            current_amount = None
            start_month = None

            for month, amount in filtered_months.items():
                year = year_plus1 if month in ["July", "August", "September", "October", "November", "December"] else year_plus2

                if current_amount is None:
                    current_amount = amount
                    start_month = month
                elif abs(current_amount - amount) > 0.3:  # Different enough to create a new range
                    end_year = year_plus1 if start_month in ["July", "August", "September", "October", "November", "December"] else year_plus2
                    amount_ranges.append((start_month, month, current_amount, end_year))
                    current_amount = amount
                    start_month = month
                # Handle the last range
                if month == "June":
                    end_year = year_plus2 
                    amount_ranges.append((start_month, month, current_amount, end_year))

            for start, end, amount, end_year in amount_ranges:
                para.add_run(f'${amount:,.2f}/month from {start} {year_plus1} to {end} {end_year}\n')

# Section: GST Credit
def gst_credit(doc, return_summary, isNewcomer, year, cfg=None):
    para = doc.add_paragraph()
    gst_title_text = get_text(cfg, 'gstTitle', 'GST/HST Credits:')
    title_run = para.add_run(f'{gst_title_text} ')
    apply_style(title_run, get_style(cfg, 'gstTitle') or {'bold': True, 'underline': True})
    if isNewcomer:
      title_run = para.add_run('* ')
      title_run.bold = True
      title_run.underline = True
    title_run = para.add_run(':\n')

    title_run.bold = True
    title_run.underline = True

    gst_credit_amount = return_summary["gst_amounts"]["gst_credit_amount"]
    july_amount = return_summary["gst_amounts"].get("july_amount", 0)
    october_amount = return_summary["gst_amounts"].get("october_amount", 0)
    january_amount = return_summary["gst_amounts"].get("january_amount", 0)
    april_amount = return_summary["gst_amounts"].get("april_amount", 0)

    para.add_run(f'You will receive a total of ').bold = False
    total_run = para.add_run(f"${gst_credit_amount:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' for the ')
    para.add_run(f'GST/ HST credit ').bold = True
    para.add_run(f'as follows:\n')

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2
    
    if july_amount > 0:
        para.add_run(f'July {year_plus1}: ${july_amount:,.2f}\n')
    if october_amount > 0:
        para.add_run(f'October {year_plus1}: ${october_amount:,.2f}\n')
    if january_amount > 0:
        para.add_run(f'January {year_plus2}: ${january_amount:,.2f}\n')
    if april_amount > 0:
        para.add_run(f'April {year_plus2}: ${april_amount:,.2f}\n')

    if isNewcomer:
        para.add_run('\n*Note that you will receive a letter from Canada Revenue Agency asking you to provide your income before arrival to Canada (so from January 1st until the date of arrival). Even though it was mentioned on the declaration, you still need to respond to the letter and provide the amount. If you do not reply, they will not pay the GST amount.\n').italic = True

from docx.shared import RGBColor

# Section: ECGEB Credit
def ecgeb_credit(doc, return_summary, isNewcomer, year, cfg=None):
    para = doc.add_paragraph()
    ecgeb_title_text = get_text(cfg, 'ecgebTitle', 'Canada Groceries & Essentials Benefit:')
    title_run = para.add_run(f'{ecgeb_title_text} ')
    apply_style(title_run, get_style(cfg, 'ecgebTitle') or {'bold': True, 'underline': True})
    if isNewcomer:
        title_run = para.add_run('* ')
        title_run.bold = True
        title_run.underline = True
    title_run = para.add_run(':\n')
    title_run.bold = True
    title_run.underline = True

    ecgeb_credit_amount = return_summary["ecgeb_amounts"]["ecgeb_credit_amount"]
    july_amount = return_summary["ecgeb_amounts"].get("july_amount", 0)
    october_amount = return_summary["ecgeb_amounts"].get("october_amount", 0)
    january_amount = return_summary["ecgeb_amounts"].get("january_amount", 0)
    april_amount = return_summary["ecgeb_amounts"].get("april_amount", 0)

    para.add_run(f'You will receive a total of ').bold = False
    total_run = para.add_run(f"${ecgeb_credit_amount:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' for the ')
    para.add_run(f'Canada Groceries and Essentials Benefit (ECGEB) ').bold = True
    para.add_run(f'as follows:\n')

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    if july_amount > 0:
        para.add_run(f'July {year_plus1}: ${july_amount:,.2f}\n')
    if october_amount > 0:
        para.add_run(f'October {year_plus1}: ${october_amount:,.2f}\n')
    if january_amount > 0:
        para.add_run(f'January {year_plus2}: ${january_amount:,.2f}\n')
    if april_amount > 0:
        para.add_run(f'April {year_plus2}: ${april_amount:,.2f}\n')

    if isNewcomer:
        para.add_run('\n*Note that you will receive a letter from Canada Revenue Agency asking you to provide your income before arrival to Canada (so from January 1st until the date of arrival). Even though it was mentioned on the declaration, you still need to respond to the letter and provide the amount. If you do not reply, they will not pay the ECGEB amount.\n').italic = True

# Section: Carbon Rebate
def carbon_rebate(doc, return_summary,  year, cfg=None):
    para = doc.add_paragraph()
    carbon_title_text = get_text(cfg, 'carbonRebateTitle', 'Canada Carbon Rebate:')
    title_run = para.add_run(f'{carbon_title_text} ')
    apply_style(title_run, get_style(cfg, 'carbonRebateTitle') or {'bold': True, 'underline': True})

    title_run = para.add_run('Credits:\n')

    title_run.bold = True
    title_run.underline = True

    carbon_rebate_amount = return_summary["carbon_rebate_amounts"]["carbon_rebate_amount"]
    july_amount = return_summary["carbon_rebate_amounts"].get("july_amount", 0)
    october_amount = return_summary["carbon_rebate_amounts"].get("october_amount", 0)
    january_amount = return_summary["carbon_rebate_amounts"].get("january_amount", 0)
    april_amount = return_summary["carbon_rebate_amounts"].get("april_amount", 0)

    para.add_run(f'You will receive a total of ').bold = False
    total_run = para.add_run(f"${carbon_rebate_amount:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' for the ')
    para.add_run(f'Canada Carbon Rebate ').bold = True
    para.add_run(f'as follows:\n')

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2
    
    # Updated months order
    if april_amount > 0:
        para.add_run(f'April {year_plus1}: ${april_amount:,.2f}\n')
    if july_amount > 0:
        para.add_run(f'July {year_plus1}: ${july_amount:,.2f}\n')
    if october_amount > 0:
        para.add_run(f'October {year_plus1}: ${october_amount:,.2f}\n')
    if january_amount > 0:
        para.add_run(f'January {year_plus2}: ${january_amount:,.2f}\n')

# Section: Climate Action Credit
def climate_action_credit(doc, return_summary, year, cfg=None):
    para = doc.add_paragraph()
    climate_title_text = get_text(cfg, 'climateActionTitle', 'BC Climate Action Tax Credit:')
    title_run = para.add_run(f'{climate_title_text} ')
    apply_style(title_run, get_style(cfg, 'climateActionTitle') or {'bold': True, 'underline': True})

    title_run = para.add_run(':\n')
    title_run.bold = True
    title_run.underline = True

    climate_action_amount = return_summary["climate_action_credit_amounts"]["climate_action_amount"]
    july_amount = return_summary["climate_action_credit_amounts"].get("july_amount", 0)
    october_amount = return_summary["climate_action_credit_amounts"].get("october_amount", 0)
    january_amount = return_summary["climate_action_credit_amounts"].get("january_amount", 0)
    april_amount = return_summary["climate_action_credit_amounts"].get("april_amount", 0)

    para.add_run('You will receive a total of ')
    total_run = para.add_run(f"${climate_action_amount:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(' for the ')
    para.add_run('Climate Action Credit ').bold = True
    para.add_run('as follows:\n')

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    # Print in chronological order: July {year+1}, October {year+1}, January {year+2}, April {year+2}
    if july_amount > 0:
        para.add_run(f"July {year_plus1}: ${july_amount:,.2f}\n")
    if october_amount > 0:
        para.add_run(f"October {year_plus1}: ${october_amount:,.2f}\n")
    if january_amount > 0:
        para.add_run(f"January {year_plus2}: ${january_amount:,.2f}\n")
    if april_amount > 0:
        para.add_run(f"April {year_plus2}: ${april_amount:,.2f}\n")

    return doc

# Section: Ontario Trillium Benefit
def ontario_trillium_benefit(doc, ontario_trillium_result, year, cfg=None):
    para = doc.add_paragraph()
    # Add title directly with bold and underline
    ot_title = get_text(cfg, 'ontarioTrilliumTitle', 'Ontario Trillium Benefit:')
    title_run = para.add_run(f'{ot_title}\n')
    apply_style(title_run, get_style(cfg, 'ontarioTrilliumTitle') or {'bold': True, 'underline': True})

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    ontario_trillium_amount = ontario_trillium_result["ontario_trillium_amount"]

    para.add_run(f'You will receive a total of ').bold = False
    total_run = para.add_run(f"${ontario_trillium_amount:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' for the ')
    para.add_run(f'Ontario Trillium Benefit ').bold = True
    para.add_run(f'as follows:\n')

    # Extract amounts for each month
    amounts = {
        "July": ontario_trillium_result["july_amount"],
        "August": ontario_trillium_result["august_amount"],
        "September": ontario_trillium_result["september_amount"],
        "October": ontario_trillium_result["october_amount"],
        "November": ontario_trillium_result["november_amount"],
        "December": ontario_trillium_result["december_amount"],
        "January": ontario_trillium_result["january_amount"],
        "February": ontario_trillium_result["february_amount"],
        "March": ontario_trillium_result["march_amount"],
        "April": ontario_trillium_result["april_amount"],
        "May": ontario_trillium_result["may_amount"],
        "June": ontario_trillium_result["june_amount"],
    }

    # Filter out months with a zero amount
    filtered_months = {month: amount for month, amount in amounts.items() if amount > 0}

    if len(filtered_months) <= 4:
        for month, amount in filtered_months.items():
            # Adjust year based on month and year+1 or year+2
            if month in ["July", "August", "September", "October", "November", "December"]:
                para.add_run(f'{month} {year_plus1}: ${amount:,.2f}\n')
            else:
                para.add_run(f'{month} {year_plus2}: ${amount:,.2f}\n')
    else:
        # Check for consistency of amounts for the range
        unique_amounts = set(filtered_months.values())

        if len(unique_amounts) == 1:
            # If all amounts are the same
            amount = next(iter(unique_amounts))  # Get the single value
            para.add_run(f'${amount:,.2f}/month from July {year_plus1} to June {year_plus2}\n')
        else:
            # Group amounts by ranges
            amount_ranges = []
            current_amount = None
            start_month = None

            for month, amount in filtered_months.items():
                # Adjust year based on month
                if month in ["July", "August", "September", "October", "November", "December"]:
                    current_year = year_plus1
                else:
                    current_year = year_plus2

                if current_amount is None:
                    current_amount = amount
                    start_month = month
                elif abs(current_amount - amount) > 0.3:  # Different enough to create a new range
                    amount_ranges.append((start_month, month, current_amount, current_year))
                    current_amount = amount
                    start_month = month

                # Handle the last range
                if month == "June":
                    amount_ranges.append((start_month, month, current_amount, current_year))

            for start, end, amount, end_year in amount_ranges:
                para.add_run(f'${amount:,.2f}/month from {start} {year_plus1} to {end} {end_year}\n')



    
# Section: Summary of Carryforward Amounts
def carryforward_amounts(doc, return_summary, cfg=None):
    para = doc.add_paragraph()
    tuition_title = get_text(cfg, 'tuitionTitle', 'Your accumulated tuition fees carried forward to future years:')
    title_run = para.add_run(f'{tuition_title}\n')
    apply_style(title_run, get_style(cfg, 'tuitionTitle') or {'bold': True})

    federal_tuition_amount = return_summary["carryforward_amounts"]["federal_tuition_amount"]
    quebec_tuition_8_percent = return_summary["carryforward_amounts"]["quebec_tuition_8_percent"]

    # Convert string values safely (removing commas for conversion)
    if isinstance(federal_tuition_amount, str):
        if federal_tuition_amount.strip():
            federal_tuition_amount = int(federal_tuition_amount.replace(",", ""))
        else:
            federal_tuition_amount = 0
    if isinstance(quebec_tuition_8_percent, str):
        if quebec_tuition_8_percent.strip():
            quebec_tuition_8_percent = int(quebec_tuition_8_percent.replace(",", ""))
        else:
            quebec_tuition_8_percent = 0

    # Format numbers with commas as thousand separators when adding them to the document
    if federal_tuition_amount > 0:
        fed_label = get_text(cfg, 'tuitionFedLabel', 'Federal (eligible to 15%):  $')
        para.add_run(fed_label)
        para.add_run(f"{federal_tuition_amount:,}")
    if 'quebec_tuition_8_percent' in locals() and quebec_tuition_8_percent is not None and quebec_tuition_8_percent > 0:
        qc_label = get_text(cfg, 'tuitionQCLabel', 'QC (eligible to 8%): $')
        para.add_run(f'\n{qc_label}')
        para.add_run(f"{quebec_tuition_8_percent:,}")

    tuition_exp = get_text(cfg, 'tuitionExplanation', 'Those accumulated tuition fees are tax credits that you will be using in future tax declarations when you work and pay tax on your income.')
    r = para.add_run(f'\n\n{tuition_exp}')
    apply_style(r, get_style(cfg, 'tuitionExplanation') or {'italic': True})




# Section: Child Benefit Amount
def child_benefit(doc, return_summary, year, cfg=None):
    para = doc.add_paragraph()
    # Add title directly with bold and underline
    ccb_title = get_text(cfg, 'ccbTitle', 'Canada Child Benefit (CCB):')
    title_run = para.add_run(f'{ccb_title}\n')
    apply_style(title_run, get_style(cfg, 'ccbTitle') or {'bold': True, 'underline': True})

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    # Extract total Canada Child Benefit amount
    total_ccb = return_summary["ccb_amounts"]["ccb_amount"]
    amounts = {
        f"July {year_plus1}": return_summary["ccb_amounts"]["july_amount"],
        f"August {year_plus1}": return_summary["ccb_amounts"]["august_amount"],
        f"September {year_plus1}": return_summary["ccb_amounts"]["september_amount"],
        f"October {year_plus1}": return_summary["ccb_amounts"]["october_amount"],
        f"November {year_plus1}": return_summary["ccb_amounts"]["november_amount"],
        f"December {year_plus1}": return_summary["ccb_amounts"]["december_amount"],
        f"January {year_plus2}": return_summary["ccb_amounts"]["january_amount"],
        f"February {year_plus2}": return_summary["ccb_amounts"]["february_amount"],
        f"March {year_plus2}": return_summary["ccb_amounts"]["march_amount"],
        f"April {year_plus2}": return_summary["ccb_amounts"]["april_amount"],
        f"May {year_plus2}": return_summary["ccb_amounts"]["may_amount"],
        f"June {year_plus2}": return_summary["ccb_amounts"]["june_amount"],
    }


    # Display total benefit amount
    para.add_run(f'You will receive a total of ').bold = False
    total_run = para.add_run(f"${total_ccb:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' for the ')
    para.add_run(f'Canada Child Benefit ').bold = True
    para.add_run(f'as follows:\n')

    # Filter out months with a zero amount
    filtered_months = {month: amount for month, amount in amounts.items() if amount > 0}

    if len(filtered_months) <= 4:
        for month, amount in filtered_months.items():
            para.add_run(f'{month}: ${amount:,.2f}\n')
    else:
        # Check for consistency of amounts for the range
        unique_amounts = set(filtered_months.values())

        if len(unique_amounts) == 1:
            # If all amounts are the same
            amount = next(iter(unique_amounts))  # Get the single value
            para.add_run(f'${amount:,.2f}/month from July {year_plus1} to June {year_plus2}')
        else:
            # Group amounts by ranges
            amount_ranges = []
            current_amount = None
            start_month = None
            prev_month = None

            for month, amount in filtered_months.items():
                # Automatically determine the year from month
                year = month.split()[1]  # Extract year from month string

                if current_amount is None:
                    current_amount = amount
                    start_month = month
                elif abs(current_amount - amount) > 0.3:
                    amount_ranges.append((start_month, prev_month, current_amount, year))
                    current_amount = amount
                    start_month = month

                prev_month = month

            # Handle the last range
            if start_month:
                amount_ranges.append((start_month, prev_month, current_amount, year))

            for start, end, amount, year in amount_ranges:
                if start == end:
                    para.add_run(f'${amount:,.2f} in {start}\n')
                else:
                    para.add_run(f'${amount:,.2f}/month from {start} to {end}\n')

# Section: Quebec Family Allowance
def quebec_family_allowance(doc, return_summary, year, cfg=None):
    para = doc.add_paragraph()

    total_allowance_amount = return_summary["family_allowance_amounts"]["fa_amount"]

    para.add_run(f'You will receive a total of ').bold = False
    total_run = para.add_run(f"${total_allowance_amount:,.2f}")
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' for the ')
    para.add_run(f'Quebec Family Allowance ').bold = True
    para.add_run(f'as follows:\n')

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    july_amount = return_summary["family_allowance_amounts"]["july_amount"]
    october_amount = return_summary["family_allowance_amounts"]["october_amount"]
    january_amount = return_summary["family_allowance_amounts"]["january_amount"]
    april_amount = return_summary["family_allowance_amounts"]["april_amount"]

    if july_amount > 0:
        para.add_run(f'July {year_plus1}: ${july_amount:,.2f}\n')
    if october_amount > 0:
        para.add_run(f'October {year_plus1}: ${october_amount:,.2f}\n')
    if january_amount > 0:
        para.add_run(f'January {year_plus2}: ${january_amount:,.2f}\n')
    if april_amount > 0:
        para.add_run(f'April {year_plus2}: ${april_amount:,.2f}\n')

# Section: Conclusion
def conclusion(doc, isMailQC=False, cfg=None):
    para = doc.add_paragraph()

    # Add the red text above the conclusion
    form_text = "form" if isMailQC else "forms"
    waiting_default = f'We will be waiting for the signed authorization {form_text} to proceed.'
    waiting_text = get_text(cfg, 'conclusionWaiting', waiting_default)
    red_run = para.add_run(f'\n{waiting_text}\n')
    apply_style(red_run, get_style(cfg, 'conclusionWaiting') or {'color': '#cd3350'})

    # Add normal text for "Thank you."
    thank_text = get_text(cfg, 'thankYou', 'Thank you.')
    para.add_run(f'\n{thank_text}\n')

    # Add the main conclusion text with specified formatting
    disclaimer_default = (
        'We at Sankari Inc. are pleased to respond to your tax inquiries and/or file your tax returns based on the '
        'information that you provide. Inaccurate or incomplete information provided by you may lead to inadequate or '
        'incorrect advice for which Sankari Inc. team cannot be held responsible. You, the client, are responsible for '
        'giving correct information and documentation to Sankari Inc.'
    )
    disclaimer_text = get_text(cfg, 'disclaimer', disclaimer_default)
    conclusion_run = para.add_run(f'\n{disclaimer_text}')
    apply_style(conclusion_run, get_style(cfg, 'disclaimer') or {'fontSize': 8, 'color': '#7f7f7f', 'italic': True})

def createIndividualWordDocFR(individual, output_file_path, doc_text_config=None):
    return_summary = individual['summary']
    year = individual['year']
    ind_title = individual['title']
    if individual['title'] == 'Mr':
        ind_title = "M."
    elif individual['title'] == 'Ms' or individual['title'] == 'Mrs':
        ind_title = "Mme"
    isMailQC = individual['isMailQC']
    isNewcomer = individual['isNewcomer']
    doc = Document()

    cfg = get_cfg(doc_text_config, resolve_doc_type_key('FR', is_couple=False))

    set_default_font(doc, "Calibri", 10)
    section_1FR(doc, return_summary, ind_title, year, isMailQC, cfg=cfg)
    tax_returnFR(doc, return_summary, year, cfg=cfg)

    if "gst_amounts" in return_summary and return_summary["gst_amounts"]:
        gst_creditFR(doc, return_summary, isNewcomer, year, cfg=cfg)

    if "ecgeb_amounts" in return_summary and return_summary["ecgeb_amounts"]:
        ecgeb_creditFR(doc, return_summary, isNewcomer, year, cfg=cfg)

    if "solidarity_amounts" in return_summary and return_summary["solidarity_amounts"]:
        solidarity_creditFR(doc, return_summary, year, cfg=cfg)

    if "carbon_rebate_amounts" in return_summary and return_summary["carbon_rebate_amounts"]:
        carbon_rebateFR(doc, return_summary, year, cfg=cfg)

    if "ccb_amounts" in return_summary and return_summary["ccb_amounts"]:
        child_benefitFR(doc, return_summary, year, cfg=cfg)

    if "family_allowance_amounts" in return_summary and return_summary["family_allowance_amounts"]:
        quebec_family_allowanceFR(doc, return_summary, year, cfg=cfg)

    if "carryforward_amounts" in return_summary and return_summary["carryforward_amounts"]:
        carryforward_amountsFR(doc, return_summary, cfg=cfg)

    conclusionFR(doc, isMailQC, cfg=cfg)
    # Save the document in the specified path
    doc.save(output_file_path)

def createCoupleWordDocFR(couple_summaries, output_file_path, doc_text_config=None):
    doc = Document()

    cfg = get_cfg(doc_text_config, resolve_doc_type_key('FR', is_couple=True))

    # Set default font to Calibri and size 10pt for the entire document
    set_default_font(doc, "Calibri", 10)

    # Initialize variables to track the year, isMailQC, and isNewcomer
    year = couple_summaries[0]['year']
    isMailQC = False
    isNewcomer = False

    for individual in couple_summaries:
        if individual['isMailQC']:
            isMailQC = True
        if individual['isNewcomer']:
            isNewcomer = True

    # Process the primary individual
    primary_individual = next(individual for individual in couple_summaries if individual['isPrimary'])
    secondary_individual = next(individual for individual in couple_summaries if not individual['isPrimary'])

    primary_title = primary_individual['title']
    if primary_individual['title'] == 'Mr':
        primary_title = "M."
    elif individual['title'] == 'Ms' or individual['title'] == 'Mrs':
        primary_title = "Mme"

    secondary_title = secondary_individual['title']
    if secondary_individual['title'] == 'Mr':
        secondary_title = "M."
    elif secondary_individual['title'] == 'Ms' or secondary_individual['title'] == 'Mrs':
        secondary_title = "Mme"

    section_1FR(doc, primary_individual['summary'], primary_title, year, isMailQC, couple=True, secondary_summary=secondary_individual['summary'], secondary_ind_title=secondary_title, cfg=cfg)
    tax_returnFR(doc, primary_individual['summary'], year,  secondary_individual['summary'], primary_title, secondary_title, isCouple=True, cfg=cfg),

    for individual in couple_summaries:
        return_summary = individual['summary']
        if "gst_amounts" in return_summary and return_summary["gst_amounts"]:
            gst_creditFR(doc, return_summary, isNewcomer, year, cfg=cfg)

        if "ecgeb_amounts" in return_summary and return_summary["ecgeb_amounts"]:
            ecgeb_creditFR(doc, return_summary, isNewcomer, year, cfg=cfg)

        if "solidarity_amounts" in return_summary and return_summary["solidarity_amounts"]:
            solidarity_creditFR(doc, return_summary, year, cfg=cfg)

        if "carbon_rebate_amounts" in return_summary and return_summary["carbon_rebate_amounts"]:
            carbon_rebateFR(doc, return_summary, year, cfg=cfg)


        if "ccb_amounts" in return_summary and return_summary["ccb_amounts"]:
            child_benefitFR(doc, return_summary, year, cfg=cfg)

        if "family_allowance_amounts" in return_summary and return_summary["family_allowance_amounts"]:
            quebec_family_allowanceFR(doc, return_summary, year, cfg=cfg)

    conclusionFR(doc, isMailQC=isMailQC, cfg=cfg)

    # Save the document in the specified path
    doc.save(output_file_path)

def section_1FR(doc, return_summary, ind_title, year, isMailQC, couple=False, secondary_summary=None, secondary_ind_title=None, cfg=None):
    # Title: Centered, Dark Gray, 14pt, Calibri
    if couple:
        default_title_fr = f"Sommaire de vos d\u00e9clarations d\u2019imp\u00f4ts {year}"
    else:
        default_title_fr = f"Sommaire de votre d\u00e9claration d\u2019imp\u00f4ts {year}"
    title_text = get_text(cfg, 'docTitle', default_title_fr, year=year)
    title = doc.add_paragraph(title_text)
    title_style = get_style(cfg, 'docTitle') or {'fontSize': 14, 'color': '#414141'}
    apply_alignment(title, title_style)
    if not get_style(cfg, 'docTitle'):
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    apply_style(run, title_style)

    # Recipient Name
    para = doc.add_paragraph()
    first_name = return_summary["tax_summary"]["first_name"]
    last_name = return_summary["tax_summary"]["last_name"]
    province = return_summary["tax_summary"]["province"]
    para.add_run(f"{ind_title} {last_name}")

    # Display Secondary Individual’s name if applicable
    if couple and secondary_summary is not None and secondary_ind_title is not None:
        secondary_first_name = secondary_summary["tax_summary"]["first_name"]
        secondary_last_name = secondary_summary["tax_summary"]["last_name"]
        para.add_run(f" & {secondary_ind_title} {secondary_last_name}")

    # Introduction paragraph
    para = doc.add_paragraph()
    default_attach = f'Nous avons joint \u00e0 ce courriel tous les documents de {"vos d\u00e9clarations" if couple else "votre d\u00e9claration"} d\u2019imp\u00f4ts {year}. '
    styled_run(para, cfg, 'introAttachment', default_attach, {'bold': False}, year=year)
    styled_run(para, cfg, 'introPassword', 'Le mot de passe se compose des neuf chiffres de votre num\u00e9ro d\u2019assurance sociale.\n', {'bold': True})
    if couple:
        styled_run(para, cfg, 'introCopyDescription', 'Le document nomm\u00e9 COPIE est une copie compl\u00e8te de votre d\u00e9claration de revenus. ', {'bold': False})
        styled_run(para, cfg, 'introCopyNoPrint', 'Vous n\'avez pas besoin de les imprimer ou de les signer; ', {'bold': True})
        styled_run(para, cfg, 'introCopyKeep', 'vous avez juste besoin de les retenir pour votre dossier. Veuillez revoir les d\u00e9clarations de revenu attentivement afin de vous assurer qu\u2019elles sont exactes et compl\u00e8tes.\n', {'bold': False})
    else:
        styled_run(para, cfg, 'introCopyDescription', 'Le document nomm\u00e9 COPIE est une copie compl\u00e8te de votre d\u00e9claration de revenus. ', {'bold': False})
        styled_run(para, cfg, 'introCopyNoPrint', 'Vous n\'avez pas besoin de l\u2019imprimer ou de le signer; ', {'bold': True})
        styled_run(para, cfg, 'introCopyKeep', 'vous avez juste besoin de le retenir pour votre dossier. Veuillez revoir la d\u00e9claration de revenu attentivement afin de vous assurer qu\u2019elle est exacte et compl\u00e8te.', {'bold': False})

    # "Very Important" without extra paragraph spacing
    vi_text = get_text(cfg, 'veryImportantHeading', '** TR\u00c8S IMPORTANT :')
    very_important_run = para.add_run(f'\n\n{vi_text}\n')
    apply_style(very_important_run, get_style(cfg, 'veryImportantHeading') or {'bold': True, 'underline': True, 'color': '#cd3350'})

    if province == "Qu\u00e9bec":

        if isMailQC:
            # Special message if isMailQC is True
            # Federal tax return
            fed_title = get_text(cfg, 'qcMailFedTitle', 'D\u00e9claration F\u00e9d\u00e9rale :')
            federal_title_run = para.add_run(f'{fed_title}\n')
            apply_style(federal_title_run, get_style(cfg, 'qcMailFedTitle') or {'bold': True, 'underline': True})

            fed_not_submitted = get_text(cfg, 'qcMailFedNotSubmitted', 'Notez que votre d\u00e9claration d\u2019imp\u00f4t F\u00e9d\u00e9rale n\u2019a pas encore \u00e9t\u00e9 soumise au gouvernement via EFile.')
            r = para.add_run(f'{fed_not_submitted}\n')
            apply_style(r, get_style(cfg, 'qcMailFedNotSubmitted') or {'bold': True})

            name = f"{first_name} {last_name}"
            sec_name = f"{secondary_first_name} {secondary_last_name}" if couple and secondary_summary else ""
            auth_text = get_text(cfg, 'qcMailFedAuthForm',
                f'Vous trouverez ci-joint {"deux formulaires" if couple else "un formulaire"}  d\u2019autorisation. Veuillez le signer \u00e9lectroniquement (ou imprimer/signer) et nous l\u2019envoyer par courriel le plus t\u00f4t possible afin que nous puissions transmettre votre d\u00e9claration par EFILE.')
            para.add_run(f'{auth_text}\n')

            sign_text = get_text(cfg, 'qcMailSignPartF', 'S\u2019il vous pla\u00eet signer la partie F.')
            para.add_run(f'{sign_text}\n\n')

            # Qu\u00e9bec tax return
            qc_title = get_text(cfg, 'qcMailQCTitle', 'D\u00e9claration Provinciale :')
            quebec_title_run = para.add_run(f'{qc_title}\n')
            apply_style(quebec_title_run, get_style(cfg, 'qcMailQCTitle') or {'bold': True, 'underline': True})

            cannot_efile = get_text(cfg, 'qcMailQCCannotEfile', 'Veuillez noter que votre d\u00e9claration provinciale ne peut pas \u00eatre transmise via Efile.')
            para.add_run(f'{cannot_efile}\n')

            if couple and secondary_summary is not None:
                print_text = get_text(cfg, 'qcMailQCPrint',
                    f'Pour cette raison, vous devez imprimer le document "QC {year} - {first_name} {last_name}.pdf" et "QC {year} - {secondary_first_name} {secondary_last_name}.pdf", signer en bas de la page ##, et l\u2019envoyer par la poste \u00e0 l\u2019adresse suivante :',
                    year=year, name=name, primaryName=name, secondaryName=sec_name)
            else:
                print_text = get_text(cfg, 'qcMailQCPrint',
                    f'Pour cette raison, vous devez imprimer le document "QC {year} - {first_name} {last_name}.pdf", signer en bas de la page ##, et l\u2019envoyer par la poste \u00e0 l\u2019adresse suivante :',
                    year=year, name=name, primaryName=name, secondaryName=sec_name)
            para.add_run(f'{print_text}\n')

            # Address in italics
            addr_text = get_text(cfg, 'qcAddress', 'Revenu Qu\u00e9bec\nC. P. 2500, succursale Place-Desjardins\nMontr\u00e9al (Qu\u00e9bec) H5B 1A3')
            address = para.add_run(f'{addr_text}\n\n')
            apply_style(address, get_style(cfg, 'qcAddress') or {'italic': True})

            on_behalf = get_text(cfg, 'qcMailOnBehalf',
                '*Si vous souhaitez qu\u2019on s\u2019occupe d\u2019envoyer la d\u00e9claration par la poste \u00e0 Revenu QC, veuillez nous envoyer par courriel la d\u00e9claration sign\u00e9e (par signature \u00e9lectronique qui ressemble \u00e0 votre signature) et nous l\u2019imprimerons et l\u2019enverrons par courrier recommand\u00e9 (avec un num\u00e9ro de suivi) \u00e0 Revenu QC. '
                'Veuillez noter qu\'il y aura des frais de service suppl\u00e9mentaires de 25 $ plus les frais de Postes Canada.')
            para.add_run(f'{on_behalf}\n')
        else:
            # Original message if isMailQC is False
            not_submitted = get_text(cfg, 'qcEfileNotSubmitted', 'Notez que votre d\u00e9claration d\u2019imp\u00f4t n\u2019a pas encore \u00e9t\u00e9 soumise au gouvernement via EFile.')
            r = para.add_run(f'{not_submitted}\n')
            apply_style(r, get_style(cfg, 'qcEfileNotSubmitted') or {'bold': True})

            auth_text = get_text(cfg, 'qcEfileAuthForms',
                f'Vous trouverez ci-joint {"quatre" if couple else "deux"} formulaires d\u2019autorisation. Veuillez les signer \u00e9lectroniquement (ou imprimer/signer) et nous les envoyer par courriel le plus t\u00f4t possible afin que nous puissions transmettre votre d\u00e9claration par EFILE.')
            para.add_run(f'{auth_text}\n')

            sign_fed = get_text(cfg, 'qcEfileSignFed', 'Pour le formulaire F\u00e9d\u00e9ral, veuillez signer la partie F.')
            r = para.add_run(f'{sign_fed}\n')
            apply_style(r, get_style(cfg, 'qcEfileSignFed') or {'italic': True})

            sign_qc = get_text(cfg, 'qcEfileSignQC', 'Pour le formulaire du Qu\u00e9bec, veuillez signer \u00e0 la fin de la section 4.')
            r = para.add_run(sign_qc)
            apply_style(r, get_style(cfg, 'qcEfileSignQC') or {'italic': True})
    else:
        # Non-Qu\u00e9bec clients
        not_submitted = get_text(cfg, 'nonQcNotSubmitted', 'Noter que votre d\u00e9claration d\u2019imp\u00f4t n\u2019a pas encore \u00e9t\u00e9 soumise au gouvernement via EFile.')
        r = para.add_run(f'{not_submitted}\n')
        apply_style(r, get_style(cfg, 'nonQcNotSubmitted') or {'bold': True})

        auth_text = get_text(cfg, 'nonQcAuthForm',
            f'Vous trouverez ci-joint {"deux formulaires" if couple else "une formulaire"}  d\u2019autorisation. Veuillez les signer \u00e9lectroniquement (ou imprimer/signer) et nous les envoyer par courriel le plus t\u00f4t possible afin que nous puissions transmettre votre d\u00e9claration par EFILE.')
        para.add_run(f'{auth_text}\n')

        sign_text = get_text(cfg, 'nonQcSignPartF', 'Veuillez signer la partie F.')
        r = para.add_run(f'{sign_text}\n')
        apply_style(r, get_style(cfg, 'nonQcSignPartF') or {'italic': True})

def tax_returnFR(doc, return_summary, year, secondary_summary=None, primary_title=None, secondary_title=None, isCouple=False, cfg=None):
    province = return_summary["tax_summary"]["province"]

    para = doc.add_paragraph()

    # Add "RÉSULTATS" title in bold and red
    results_text = get_text(cfg, 'resultsHeading', 'RÉSULTATS')
    bold_red_results = para.add_run(f'{results_text}\n')
    apply_style(bold_red_results, get_style(cfg, 'resultsHeading') or {'bold': True, 'color': '#cd3350'})
    
    
    # Helper function for formatting numbers in the French style
    def format_french_number(amount):
        return f"{amount:,.2f}".replace(",", " ").replace(".", ",")

    # Federal and Quebec Tax Return for primary individual
    if isCouple:
      para.add_run(f"\n{primary_title} {return_summary['tax_summary']['last_name']}\n")

    # Federal Tax Return
    fed_label = get_text(cfg, 'federalReturnLabel', 'Déclaration Fédérale')
    r = para.add_run(f'{fed_label}\n')
    apply_style(r, get_style(cfg, 'federalReturnLabel') or {'bold': True})
    federal_refund = return_summary["tax_summary"]["federal_refund"]
    federal_owing = return_summary["tax_summary"]["federal_owing"]

    if federal_refund > 2:
        refund_prefix = get_text(cfg, 'refundPrefix', 'Vous avez droit à un remboursement de')
        para.add_run(f"{refund_prefix} ").bold = False
        refund_run = para.add_run(f"{format_french_number(federal_refund)} $\n")
        refund_run.bold = True
        refund_run.font.color.rgb = RGBColor(0, 128, 0)
    elif federal_owing > 2:
        owing_prefix = get_text(cfg, 'owingPrefix', 'Vous avez un montant dû de')
        para.add_run(f"{owing_prefix} ").bold = False
        owing_run = para.add_run(f"{format_french_number(federal_owing)} $\n")
        owing_run.bold = True
        owing_run.font.color.rgb = RGBColor(255, 0, 0)
    else:
        no_bal = get_text(cfg, 'noBalance', "Vous n’avez pas de remboursement ni de montant dû.")
        para.add_run(f"{no_bal}\n")

    if province == "Québec":
    # Quebec Tax Return
        qc_label = get_text(cfg, 'quebecReturnLabel', 'Déclaration Provinciale')
        r = para.add_run(f'{qc_label}\n')
        apply_style(r, get_style(cfg, 'quebecReturnLabel') or {'bold': True})
        quebec_refund = return_summary["tax_summary"]["quebec_refund"]
        quebec_owing = return_summary["tax_summary"]["quebec_owing"]

        if quebec_refund > 2:
            refund_prefix = get_text(cfg, 'refundPrefix', 'Vous avez droit à un remboursement de')
            para.add_run(f"{refund_prefix} ").bold = False
            refund_run = para.add_run(f"{format_french_number(quebec_refund)} $\n")
            refund_run.bold = True
            refund_run.font.color.rgb = RGBColor(0, 128, 0)
        elif quebec_owing > 2:
            owing_prefix = get_text(cfg, 'owingPrefix', 'Vous avez un montant dû de')
            para.add_run(f"{owing_prefix} ").bold = False
            owing_run = para.add_run(f"{format_french_number(quebec_owing)} $\n")
            owing_run.bold = True
            owing_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            no_bal = get_text(cfg, 'noBalance', "Vous n’avez pas de remboursement ni de montant dû.")
            para.add_run(f"{no_bal}\n")

    # Add payment section for the primary individual
    quebec_owing = locals().get('quebec_owing', 0)

    if federal_owing > 2 or quebec_owing > 2:

        para = doc.add_paragraph()
        if federal_owing > 2 and quebec_owing > 2:
            owing_text = get_text(cfg, 'paymentOwingFedAndQC', 'Vous avez un montant dû sur vos déclarations Fédérale et Provinciale ;')
        elif federal_owing > 2:
            owing_text = get_text(cfg, 'paymentOwingFed', 'Vous avez un montant dû sur votre déclaration Fédérale ;')
        else:
            owing_text = get_text(cfg, 'paymentOwingQC', 'Vous avez un montant dû sur votre déclaration Provinciale ;')
        r = para.add_run(f'{owing_text} ')
        apply_style(r, get_style(cfg, 'paymentOwingFed') or {'italic': True})

        deadline = get_text(cfg, 'paymentDeadline', f'veuillez vous assurer de payer le solde dû avant le 30 avril {int(year) + 1} pour éviter de payer des intérêts.', dueYear=str(int(year) + 1))
        r = para.add_run(f'{deadline} ')
        apply_style(r, get_style(cfg, 'paymentDeadline') or {'italic': True})

        howto = get_text(cfg, 'paymentHowTo', 'Veuillez attendre quelques jours après notre transmission par EFILE pour payer votre solde. Pour plus de détails sur la façon de payer le montant dû, veuillez cliquer sur :')
        r = para.add_run(f'{howto} ')
        apply_style(r, get_style(cfg, 'paymentHowTo') or {'italic': True})

        # Add hyperlinks for payment details
        if federal_owing > 2:
            add_hyperlink(para, 'Fédéral', 'https://www.canada.ca/fr/agence-revenu/services/paiements/paiements-arc.html')
        if federal_owing > 0 and quebec_owing > 2:
            para.add_run(' et ')
        if quebec_owing > 2:
            add_hyperlink(para, 'Québec', 'https://www.revenuquebec.ca/fr/citoyens/declaration-de-revenus/payer-ou-etre-rembourse/solde-dimpot-a-payer/')
        para.add_run('\n')

    # Carryforward amounts
    if isCouple and "carryforward_amounts" in return_summary and return_summary["carryforward_amounts"]:
        para = doc.add_paragraph()
        tuition_title = get_text(cfg, 'tuitionTitle', 'Vos frais de scolarité accumulés reportés aux années futures :')
        title_run = para.add_run(f'{tuition_title}\n')
        apply_style(title_run, get_style(cfg, 'tuitionTitle') or {'bold': True})

        federal_tuition_amount = return_summary["carryforward_amounts"]["federal_tuition_amount"]
        quebec_tuition_8_percent = return_summary["carryforward_amounts"]["quebec_tuition_8_percent"]

        # Ensure amounts are numeric (float or int)
        if isinstance(federal_tuition_amount, str):
            if federal_tuition_amount.strip():  # Check if not empty
                federal_tuition_amount = int(federal_tuition_amount)
            else:
                federal_tuition_amount = 0

        if isinstance(quebec_tuition_8_percent, str):
            if quebec_tuition_8_percent.strip():  # Check if not empty
                quebec_tuition_8_percent = int(quebec_tuition_8_percent)
            else:
                quebec_tuition_8_percent = 0

        # Format the amounts without decimals and replace commas with spaces
        formatted_federal_amount = locale.format_string("%d", federal_tuition_amount, grouping=True).replace(",", " ")
        formatted_quebec_amount = locale.format_string("%d", quebec_tuition_8_percent, grouping=True).replace(",", " ")

        # Add formatted amounts to the paragraph
        if federal_tuition_amount > 0:
            para.add_run('Fédéral (admissible à 15%) : ')
            para.add_run(f"{formatted_federal_amount} $")
        if 'quebec_tuition_8_percent' in locals() and quebec_tuition_8_percent is not None and quebec_tuition_8_percent > 0:
            para.add_run('\nQC (admissible à 8%) : ')
            para.add_run(f"{formatted_quebec_amount} $")

        # Add additional explanatory text
        para.add_run(
            '\n\nCes frais de scolarité accumulés sont des crédits d\'impôt que vous allez utiliser dans vos futures déclarations, lorsque vous générez un revenu et payez des impôts.'
        ).italic = True

    # Process secondary individual if applicable
    if isCouple and secondary_summary:
        para.add_run(f"\n{secondary_title} {secondary_summary['tax_summary']['last_name']}\n")

        # Federal Tax Return for secondary individual
        fed_label = get_text(cfg, 'federalReturnLabel', 'Déclaration Fédérale')
        r = para.add_run(f'{fed_label}\n')
        apply_style(r, get_style(cfg, 'federalReturnLabel') or {'bold': True})
        federal_refund = secondary_summary["tax_summary"]["federal_refund"]
        federal_owing = secondary_summary["tax_summary"]["federal_owing"]

        if federal_refund > 2:
            refund_prefix = get_text(cfg, 'refundPrefix', 'Vous avez droit à un remboursement de')
            para.add_run(f"{refund_prefix} ").bold = False
            refund_run = para.add_run(f"{format_french_number(federal_refund)} $\n")
            refund_run.bold = True
            refund_run.font.color.rgb = RGBColor(0, 128, 0)
        elif federal_owing > 2:
            owing_prefix = get_text(cfg, 'owingPrefix', 'Vous avez un montant dû de')
            para.add_run(f"{owing_prefix} ").bold = False
            owing_run = para.add_run(f"{format_french_number(federal_owing)} $\n")
            owing_run.bold = True
            owing_run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            no_bal = get_text(cfg, 'noBalance', "Vous n’avez pas de remboursement ni de montant dû.")
            para.add_run(f"{no_bal}\n")
        if province == "Québec":
            # Quebec Tax Return for secondary individual
            qc_label = get_text(cfg, 'quebecReturnLabel', 'Déclaration Provinciale')
            r = para.add_run(f'{qc_label}\n')
            apply_style(r, get_style(cfg, 'quebecReturnLabel') or {'bold': True})
            quebec_refund = secondary_summary["tax_summary"]["quebec_refund"]
            quebec_owing = secondary_summary["tax_summary"]["quebec_owing"]

            if quebec_refund > 2:
                refund_prefix = get_text(cfg, 'refundPrefix', 'Vous avez droit à un remboursement de')
                para.add_run(f"{refund_prefix} ").bold = False
                refund_run = para.add_run(f"{format_french_number(quebec_refund)} $\n")
                refund_run.bold = True
                refund_run.font.color.rgb = RGBColor(0, 128, 0)
            elif quebec_owing > 2:
                owing_prefix = get_text(cfg, 'owingPrefix', 'Vous avez un montant dû de')
                para.add_run(f"{owing_prefix} ").bold = False
                owing_run = para.add_run(f"{format_french_number(quebec_owing)} $\n")
                owing_run.bold = True
                owing_run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                no_bal = get_text(cfg, 'noBalance', "Vous n’avez pas de remboursement ni de montant dû.")
                para.add_run(f"{no_bal}\n")

        # Add payment section for the secondary individual
        quebec_owing = locals().get('quebec_owing', 0)

        if federal_owing > 2 or quebec_owing > 2:

            para = doc.add_paragraph()
            if federal_owing > 2 and quebec_owing > 2:
                owing_text = get_text(cfg, 'paymentOwingFedAndQC', 'Vous avez un montant dû sur vos déclarations Fédérale et Provinciale ;')
            elif federal_owing > 2:
                owing_text = get_text(cfg, 'paymentOwingFed', 'Vous avez un montant dû sur votre déclaration Fédérale ;')
            else:
                owing_text = get_text(cfg, 'paymentOwingQC', 'Vous avez un montant dû sur votre déclaration Provinciale ;')
            r = para.add_run(f'{owing_text} ')
            apply_style(r, get_style(cfg, 'paymentOwingFed') or {'italic': True})

            deadline = get_text(cfg, 'paymentDeadline', f'veuillez vous assurer de payer le solde dû avant le 30 avril {int(year) + 1} pour éviter de payer des intérêts.', dueYear=str(int(year) + 1))
            r = para.add_run(f'{deadline} ')
            apply_style(r, get_style(cfg, 'paymentDeadline') or {'italic': True})

            howto = get_text(cfg, 'paymentHowTo', 'Veuillez attendre quelques jours après notre transmission par EFILE pour payer votre solde. Pour plus de détails sur la façon de payer le montant dû, veuillez cliquer sur :')
            r = para.add_run(f'{howto} ')
            apply_style(r, get_style(cfg, 'paymentHowTo') or {'italic': True})

            # Add hyperlinks for payment details
            if federal_owing > 2:
                add_hyperlink(para, 'Fédéral', 'https://www.canada.ca/fr/agence-revenu/services/paiements/paiements-arc.html')
            if federal_owing > 0 and quebec_owing > 2:
                para.add_run(' et ')
            if quebec_owing > 2:
                add_hyperlink(para, 'Québec', 'https://www.revenuquebec.ca/fr/citoyens/declaration-de-revenus/payer-ou-etre-rembourse/solde-dimpot-a-payer/')
            para.add_run('\n')

        if "carryforward_amounts" in secondary_summary and secondary_summary["carryforward_amounts"]:
            para = doc.add_paragraph()
            tuition_title = get_text(cfg, 'tuitionTitle', 'Vos frais de scolarité accumulés reportés aux années futures :')
            title_run = para.add_run(f'{tuition_title}\n')
            apply_style(title_run, get_style(cfg, 'tuitionTitle') or {'bold': True})

            federal_tuition_amount = secondary_summary["carryforward_amounts"]["federal_tuition_amount"]
            quebec_tuition_8_percent = secondary_summary["carryforward_amounts"]["quebec_tuition_8_percent"]

            # Ensure amounts are numeric (float or int)
            if isinstance(federal_tuition_amount, str):
                if federal_tuition_amount.strip():  # Check if not empty
                    federal_tuition_amount = int(federal_tuition_amount)
                else:
                    federal_tuition_amount = 0

            if isinstance(quebec_tuition_8_percent, str):
                if quebec_tuition_8_percent.strip():  # Check if not empty
                    quebec_tuition_8_percent = int(quebec_tuition_8_percent)
                else:
                    quebec_tuition_8_percent = 0

            # Format the amounts without decimals and replace commas with spaces
            formatted_federal_amount = locale.format_string("%d", federal_tuition_amount, grouping=True).replace(",", " ")
            formatted_quebec_amount = locale.format_string("%d", quebec_tuition_8_percent, grouping=True).replace(",", " ")

            # Add formatted amounts to the paragraph
            if federal_tuition_amount > 0:
                para.add_run('Fédéral (admissible à 15%) : ')
                para.add_run(f"{formatted_federal_amount} $")
            if 'quebec_tuition_8_percent' in locals() and quebec_tuition_8_percent is not None and quebec_tuition_8_percent > 0:
                para.add_run('\nQC (admissible à 8%) : ')
                para.add_run(f"{formatted_quebec_amount} $")

            # Add additional explanatory text
            para.add_run(
                '\n\nCes frais de scolarité accumulés sont des crédits d\'impôt que vous allez utiliser dans vos futures déclarations, lorsque vous générez un revenu et payez des impôts.'
            ).italic = True

def solidarity_creditFR(doc, return_summary, year, cfg=None):
    para = doc.add_paragraph()

    # Add title with bold and underline
    sol_title = get_text(cfg, 'solidarityTitle', 'Crédit de solidarité :')
    title_run = para.add_run(f'{sol_title}\n')
    apply_style(title_run, get_style(cfg, 'solidarityTitle') or {'bold': True, 'underline': True})

    year_plus1 = int(year )+ 1
    year_plus2 = int(year) + 2  

    # Extract the total solidarity credit amount
    solidarity_amount = return_summary["solidarity_amounts"]["solidarity_credit_amount"]

    # Add the formatted total amount to the paragraph
    para.add_run(f'Vous allez recevoir un total de ').bold = False
    total_run = para.add_run(format_currency(solidarity_amount))
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' pour le crédit de ').bold = False
    para.add_run(f'Solidarité ').bold = True
    para.add_run(f'de la façon suivante :\n')

    # Extract amounts for each month
    amounts = {
        "Juillet": return_summary["solidarity_amounts"]["july_amount"],
        "Août": return_summary["solidarity_amounts"]["august_amount"],
        "Septembre": return_summary["solidarity_amounts"]["september_amount"],
        "Octobre": return_summary["solidarity_amounts"]["october_amount"],
        "Novembre": return_summary["solidarity_amounts"]["november_amount"],
        "Décembre": return_summary["solidarity_amounts"]["december_amount"],
        "Janvier": return_summary["solidarity_amounts"]["january_amount"],
        "Février": return_summary["solidarity_amounts"]["february_amount"],
        "Mars": return_summary["solidarity_amounts"]["march_amount"],
        "Avril": return_summary["solidarity_amounts"]["april_amount"],
        "Mai": return_summary["solidarity_amounts"]["may_amount"],
        "Juin": return_summary["solidarity_amounts"]["june_amount"],
    }

    # Filter out months with zero amounts
    filtered_months = {month: amount for month, amount in amounts.items() if amount > 0}

    # If there are 4 or fewer months with non-zero amounts
    if len(filtered_months) <= 4:
        for month, amount in filtered_months.items():
            year = year_plus1 if month in ["Juillet", "Août", "Septembre", "Octobre", "Novembre", "Décembre"] else year_plus2
            para.add_run(f'{month} {year} : {format_currency(amount)}\n')
    else:
        # Check for consistency of amounts across the range
        unique_amounts = set(filtered_months.values())

        if len(unique_amounts) == 1:
            # All amounts are the same, show it for the full period
            amount = next(iter(unique_amounts))  # Get the single unique amount
            para.add_run(f'{format_currency(amount)} /mois de Juillet {year_plus1} à Juin {year_plus2}\n')
        else:
            # Group amounts by ranges if they differ
            amount_ranges = []
            current_amount = None
            start_month = None

            for month, amount in filtered_months.items():
                year = year_plus1 if month in ["Juillet", "Août", "Septembre", "Octobre", "Novembre", "Décembre"] else year_plus2
                if current_amount is None:
                    current_amount = amount
                    start_month = month
                elif abs(current_amount - amount) > 0.3:  # Change detected, create a new range
                    amount_ranges.append((start_month, month, current_amount, year))
                    current_amount = amount
                    start_month = month

                # Handle the last range (up to June)
                if month == "Juin":
                    amount_ranges.append((start_month, month, current_amount, year_plus2))

            # Display each range
            for start, end, amount, end_year in amount_ranges:
                para.add_run(f'{format_currency(amount)} /mois de {start} {year_plus1} à {end} {end_year}\n')


def gst_creditFR(doc, return_summary, isNewcomer, year, cfg=None):
    para = doc.add_paragraph()

    # Add title with bold and underline
    gst_title_text = get_text(cfg, 'gstTitle', 'Crédits TPS/TVH :')
    title_run = para.add_run(gst_title_text)
    apply_style(title_run, get_style(cfg, 'gstTitle') or {'bold': True, 'underline': True})

    if isNewcomer:
        title_run = para.add_run('* ')
        title_run.bold = True
        title_run.underline = True

    # Continue with title
    title_run = para.add_run(' :\n')
    title_run.bold = True
    title_run.underline = True

    # Extract GST credit amounts
    gst_credit_amount = return_summary["gst_amounts"]["gst_credit_amount"]
    july_amount = return_summary["gst_amounts"].get("july_amount", 0)
    october_amount = return_summary["gst_amounts"].get("october_amount", 0)
    january_amount = return_summary["gst_amounts"].get("january_amount", 0)
    april_amount = return_summary["gst_amounts"].get("april_amount", 0)

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    # Format amounts using the helper function
    formatted_total_gst = format_currency(gst_credit_amount)

    # Add formatted total amount to the paragraph
    para.add_run(f'Vous allez recevoir un total de ').bold = False
    total_run = para.add_run(formatted_total_gst)
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' pour les crédits ')
    para.add_run(f'TPS ').bold = True
    para.add_run(f'de la façon suivante :\n')

    # Add the amounts for each month, only if greater than 0
    if july_amount > 0:
        para.add_run(f'Juillet {year_plus1} : {format_currency(july_amount)}\n')
    if october_amount > 0:
        para.add_run(f'Octobre {year_plus1} : {format_currency(october_amount)}\n')
    if january_amount > 0:
        para.add_run(f'Janvier {year_plus2} : {format_currency(january_amount)}\n')
    if april_amount > 0:
        para.add_run(f'Avril {year_plus2} : {format_currency(april_amount)}\n')

    # Newcomer note
    if isNewcomer:
        para.add_run(
            '\n*Notez que vous allez recevoir une lettre de l\'Agence du Revenu du Canada vous demandant de fournir vos revenus avant votre arrivée au Canada (du 1er janvier jusqu\'à la date d\'arrivée). '
            'Même si cela a été mentionné sur la déclaration, vous devez quand même répondre à la lettre. Si vous ne répondez pas, ils ne paieront pas le montant de la TPS.\n'
        ).italic = True


def ecgeb_creditFR(doc, return_summary, isNewcomer, year, cfg=None):
    para = doc.add_paragraph()

    ecgeb_title_text = get_text(cfg, 'ecgebTitle', "Allocation canadienne pour l\u2019\u00e9picerie :")
    title_run = para.add_run(ecgeb_title_text)
    apply_style(title_run, get_style(cfg, 'ecgebTitle') or {'bold': True, 'underline': True})

    if isNewcomer:
        title_run = para.add_run('* ')
        title_run.bold = True
        title_run.underline = True

    title_run = para.add_run(' :\n')
    title_run.bold = True
    title_run.underline = True

    ecgeb_credit_amount = return_summary["ecgeb_amounts"]["ecgeb_credit_amount"]
    july_amount = return_summary["ecgeb_amounts"].get("july_amount", 0)
    october_amount = return_summary["ecgeb_amounts"].get("october_amount", 0)
    january_amount = return_summary["ecgeb_amounts"].get("january_amount", 0)
    april_amount = return_summary["ecgeb_amounts"].get("april_amount", 0)

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    formatted_total = format_currency(ecgeb_credit_amount)

    para.add_run(f'Vous allez recevoir un total de ').bold = False
    total_run = para.add_run(formatted_total)
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' pour la ')
    para.add_run(f'Allocation canadienne pour l\'épicerie et les besoins essentiels (ACEBE) ').bold = True
    para.add_run(f'de la façon suivante :\n')

    if july_amount > 0:
        para.add_run(f'Juillet {year_plus1} : {format_currency(july_amount)}\n')
    if october_amount > 0:
        para.add_run(f'Octobre {year_plus1} : {format_currency(october_amount)}\n')
    if january_amount > 0:
        para.add_run(f'Janvier {year_plus2} : {format_currency(january_amount)}\n')
    if april_amount > 0:
        para.add_run(f'Avril {year_plus2} : {format_currency(april_amount)}\n')

    if isNewcomer:
        para.add_run(
            '\n*Notez que vous allez recevoir une lettre de l\'Agence du Revenu du Canada vous demandant de fournir vos revenus avant votre arrivée au Canada (du 1er janvier jusqu\'à la date d\'arrivée). '
            'Même si cela a été mentionné sur la déclaration, vous devez quand même répondre à la lettre. Si vous ne répondez pas, ils ne paieront pas le montant de la ACEBE.\n'
        ).italic = True


def carryforward_amountsFR(doc, return_summary, cfg=None):
    para = doc.add_paragraph()
    tuition_title = get_text(cfg, 'tuitionTitle', 'Vos frais de scolarité accumulés reportés aux années futures :')
    title_run = para.add_run(f'{tuition_title}\n')
    apply_style(title_run, get_style(cfg, 'tuitionTitle') or {'bold': True})

    federal_tuition_amount = return_summary["carryforward_amounts"]["federal_tuition_amount"]
    quebec_tuition_8_percent = return_summary["carryforward_amounts"]["quebec_tuition_8_percent"]

    # Ensure amounts are numeric (float or int)
    if isinstance(federal_tuition_amount, str):
        if federal_tuition_amount.strip():  # Check if not empty
            federal_tuition_amount = int(federal_tuition_amount)
        else:
            federal_tuition_amount = 0

    if isinstance(quebec_tuition_8_percent, str):
        if quebec_tuition_8_percent.strip():  # Check if not empty
            quebec_tuition_8_percent = int(quebec_tuition_8_percent)
        else:
            quebec_tuition_8_percent = 0

    # Format the amounts without decimals and replace commas with spaces
    formatted_federal_amount = locale.format_string("%d", federal_tuition_amount, grouping=True).replace(",", " ")
    formatted_quebec_amount = locale.format_string("%d", quebec_tuition_8_percent, grouping=True).replace(",", " ")

    # Add formatted amounts to the paragraph
    if federal_tuition_amount > 0:
        para.add_run('Fédéral (admissible à 15%) : ')
        para.add_run(f"{formatted_federal_amount} $")
    if 'quebec_tuition_8_percent' in locals() and quebec_tuition_8_percent is not None and quebec_tuition_8_percent > 0:
        para.add_run('\nQC (admissible à 8%) : ')
        para.add_run(f"{formatted_quebec_amount} $")

    # Add additional explanatory text
    para.add_run(
        '\n\nCes frais de scolarité accumulés sont des crédits d\'impôt que vous allez utiliser dans vos futures déclarations, lorsque vous générez un revenu et payez des impôts.'
    ).italic = True

def carbon_rebateFR(doc, return_summary, year, cfg=None):
    para = doc.add_paragraph()
    carbon_title_text = get_text(cfg, 'carbonRebateTitle', 'Remise canadienne sur le carbone :')
    title_run = para.add_run(f'{carbon_title_text}\n')
    apply_style(title_run, get_style(cfg, 'carbonRebateTitle') or {'bold': True, 'underline': True})

    carbon_rebate_amount = return_summary["carbon_rebate_amounts"]["carbon_rebate_amount"]
    july_amount = return_summary["carbon_rebate_amounts"].get("july_amount", 0)
    october_amount = return_summary["carbon_rebate_amounts"].get("october_amount", 0)
    january_amount = return_summary["carbon_rebate_amounts"].get("january_amount", 0)
    april_amount = return_summary["carbon_rebate_amounts"].get("april_amount", 0)

    para.add_run(f'Vous recevrez un total de ').bold = False
    total_run = para.add_run(f"${carbon_rebate_amount:,.2f}".replace(',', ' ').replace('.', ','))
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' pour la ')
    para.add_run(f'Remise canadienne sur le carbone ').bold = True
    para.add_run(f'distribuée comme suit :\n')

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    # Updated months order
    if april_amount > 0:
        para.add_run(f'Avril {year_plus1} : ${format(april_amount, ",.2f").replace(",", " ").replace(".", ",")}\n')
    if july_amount > 0:
        para.add_run(f'Juillet {year_plus1} : ${format(july_amount, ",.2f").replace(",", " ").replace(".", ",")}\n')
    if october_amount > 0:
        para.add_run(f'Octobre {year_plus1} : ${format(october_amount, ",.2f").replace(",", " ").replace(".", ",")}\n')
    if january_amount > 0:
        para.add_run(f'Janvier {year_plus2} : ${format(january_amount, ",.2f").replace(",", " ").replace(".", ",")}\n')



def child_benefitFR(doc, return_summary, year, cfg=None):
    para = doc.add_paragraph()
    # Add title directly with bold and underline
    ccb_title = get_text(cfg, 'ccbTitle', 'Allocation canadienne pour enfants (ACE) :')
    title_run = para.add_run(f'{ccb_title}\n')
    apply_style(title_run, get_style(cfg, 'ccbTitle') or {'bold': True, 'underline': True})

    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    # Extract total Canada Child Benefit amount
    total_ccb = return_summary["ccb_amounts"]["ccb_amount"]
    amounts = {
        f"Juillet {year_plus1}": return_summary["ccb_amounts"]["july_amount"],
        f"Août {year_plus1}": return_summary["ccb_amounts"]["august_amount"],
        f"Septembre {year_plus1}": return_summary["ccb_amounts"]["september_amount"],
        f"Octobre {year_plus1}": return_summary["ccb_amounts"]["october_amount"],
        f"Novembre {year_plus1}": return_summary["ccb_amounts"]["november_amount"],
        f"Décembre {year_plus1}": return_summary["ccb_amounts"]["december_amount"],
        f"Janvier {year_plus2}": return_summary["ccb_amounts"]["january_amount"],
        f"Février {year_plus2}": return_summary["ccb_amounts"]["february_amount"],
        f"Mars {year_plus2}": return_summary["ccb_amounts"]["march_amount"],
        f"Avril {year_plus2}": return_summary["ccb_amounts"]["april_amount"],
        f"Mai {year_plus2}": return_summary["ccb_amounts"]["may_amount"],
        f"Juin {year_plus2}": return_summary["ccb_amounts"]["june_amount"],
    }

    # Display total benefit amount with French number formatting
    para.add_run(f'Vous allez recevoir un total de ').bold = False
    total_run = para.add_run(format_currency(total_ccb))
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' pour ')
    para.add_run(f'l’Allocation canadienne pour enfants ').bold = True
    para.add_run(f'de la façon suivante :\n')

    # Filter out months with a zero amount
    filtered_months = {month: amount for month, amount in amounts.items() if amount > 0}

    if len(filtered_months) <= 4:
        for month, amount in filtered_months.items():
            para.add_run(f'{month} : {format_currency(amount)}\n')
    else:
        # Check for consistency of amounts for the range
        unique_amounts = set(filtered_months.values())

        if len(unique_amounts) == 1:
            # If all amounts are the same
            amount = next(iter(unique_amounts))  # Get the single value
            para.add_run(f'{format_currency(amount)} /mois de Juillet {year_plus1} à Juin {year_plus2}\n')
        else:
            # Group amounts by ranges
            amount_ranges = []
            current_amount = None
            start_month = None
            prev_month = None

            for month, amount in filtered_months.items():
                # Automatically determine the year from the month
                year = month.split()[1]  # Extract year from month string

                if current_amount is None:
                    current_amount = amount
                    start_month = month
                elif abs(current_amount - amount) > 0.3:
                    amount_ranges.append((start_month, prev_month, current_amount, year))
                    current_amount = amount
                    start_month = month

                prev_month = month  # Track the previous month

            # Handle the last range
            if start_month:
                amount_ranges.append((start_month, prev_month, current_amount, year))

            for start, end, amount, year in amount_ranges:
                if start == end:
                    para.add_run(f'{format_currency(amount)} en {start}\n')
                else:
                    para.add_run(f'{format_currency(amount)} /mois de {start} à {end}\n')


def quebec_family_allowanceFR(doc, return_summary, year, cfg=None):
    para = doc.add_paragraph()

    total_allowance_amount = return_summary["family_allowance_amounts"]["fa_amount"]

    para.add_run(f'Vous allez recevoir un total de ').bold = False
    # Format the total allowance amount with French number formatting (thousands space, comma for decimal)
    total_run = para.add_run(format_currency(total_allowance_amount))
    total_run.bold = True
    total_run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(f' pour ')
    para.add_run(f'l’Allocation Famille ').bold = True
    para.add_run(f'de la façon suivante :\n')
    
    year_plus1 = int(year) + 1
    year_plus2 = int(year) + 2

    july_amount = return_summary["family_allowance_amounts"]["july_amount"]
    october_amount = return_summary["family_allowance_amounts"]["october_amount"]
    january_amount = return_summary["family_allowance_amounts"]["january_amount"]
    april_amount = return_summary["family_allowance_amounts"]["april_amount"]

    # Format the amounts for each month if greater than 0
    if july_amount > 0:
        para.add_run(f'Juillet {year_plus1} : {format_currency(july_amount)}\n')
    if october_amount > 0:
        para.add_run(f'Octobre {year_plus1} : {format_currency(october_amount)}\n')
    if january_amount > 0:
        para.add_run(f'Janvier {year_plus2} : {format_currency(january_amount)}\n')
    if april_amount > 0:
        para.add_run(f'Avril {year_plus2} : {format_currency(april_amount)}\n')


def conclusionFR(doc, isMailQC=False, cfg=None):
    para = doc.add_paragraph()

    # Add the red text above the conclusion
    if isMailQC:
        red_run = para.add_run('\nNous attendons le formulaire d’autorisation signés pour soumettre vos déclarations.\n')
    else:
        red_run = para.add_run('\nNous attendons les formulaires d’autorisation signés pour soumettre votre déclaration.\n')
    apply_style(red_run, get_style(cfg, 'conclusionWaiting') or {'color': '#cd3350'})

    # Add normal text for "Thank you."
    thank_text = get_text(cfg, 'thankYou', 'Merci')
    para.add_run(f'\n{thank_text}\n')

    # Add the main conclusion text with specified formatting
    disclaimer_default = (
        'Nous, à Sankari Inc., sommes heureux de répondre à vos demandes de renseignements fiscaux et / ou de produire vos déclarations de revenus en '
        'fonction des renseignements que vous nous fournissez. Les informations inexactes ou incomplètes fournies par vous peuvent conduire à des '
        'conseils inadéquats ou incorrects pour lesquels l\'équipe de Sankari Inc. ne peut être tenue pour responsable. Vous, le client est responsable de '
        'fournir l\'information et la documentation correcte à l\'équipe de Sankari Inc.'
    )
    disclaimer_text = get_text(cfg, 'disclaimer', disclaimer_default)
    conclusion_run = para.add_run(f'\n{disclaimer_text}')
    apply_style(conclusion_run, get_style(cfg, 'disclaimer') or {'fontSize': 8, 'color': '#7f7f7f', 'italic': True})

