import sys
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _parse_rate(value):
    """Parse a tax rate value into a float. Accepts numbers, numeric strings, or strings with a trailing '%'."""
    try:
        if value is None:
            return 0.0
        # Convert to string and strip whitespace
        s = str(value).strip()
        if s.endswith('%'):
            s = s[:-1].strip()
        return float(s)
    except Exception:
        return 0.0

def create_confirmation_email(directory_path, clients, selected_prices, tax_rate, includeTaxes, language):

      if language == "fr":
          create_confirmation_email_french(directory_path, clients, selected_prices, tax_rate, includeTaxes)
      elif language == "en":
          create_confirmation_email_english(directory_path, clients, selected_prices, tax_rate, includeTaxes)
      else:
          raise ValueError(f"Unsupported language: {language}")

def create_confirmation_email_english(directory_path, clients, selected_prices, tax_rate, includeTaxes):
    province = tax_rate.get("province")
    # Ensure rates are numeric (allow strings like "5" or "5%")
    gst_rate = _parse_rate(tax_rate.get("fedRate", 0))
    qst_rate = _parse_rate(tax_rate.get("provRate", 0)) if province == "QC" else 0

    doc = Document()

    # Set default font to Calibri and size 10pt for the entire document
    set_default_font(doc, "Calibri", 10)

    add_intro_section(doc, clients)
    add_confirmation_numbers_section(doc, clients, province)
    if selected_prices:
      add_tax_summary_fees_section(doc, selected_prices, gst_rate, qst_rate, province, includeTaxes)
    add_footer_section(doc, province)

    years = sorted(set(year['year'] for client in clients for year in client['years']))
    if len(years) == 2:
        year_str = f"{years[0]} & {years[1]}"
    elif len(years) > 2:
        year_str = ", ".join(str(year) for year in years[:-1]) + f", & {years[-1]}"
    else:
        year_str = str(years[0])
    client_names = " & ".join([f"{client['name']}" for client in clients])
    output_file_path = f"{directory_path}/Confirmation {client_names} {year_str}.docx"

    doc.save(output_file_path)

# Function to set default font and size for the entire document
def set_default_font(doc, font_name, font_size):
    styles = doc.styles['Normal']
    font = styles.font
    font.name = font_name
    font.size = Pt(font_size)

def add_intro_section(doc, clients):
    # Get unique years from all clients
    years = sorted(set(year['year'] for client in clients for year in client['years']))

    # Format the years string
    if len(years) == 2:
        year_str = f"{years[0]} & {years[1]}"
    elif len(years) > 2:
        year_str = ", ".join(str(year) for year in years[:-1]) + f", & {years[-1]}"
    else:
        year_str = str(years[0])

    # Format the client introduction based on the number of clients
    if len(clients) == 1:
        intro = f"{clients[0]['title']} {clients[0]['name']}"
    elif len(clients) == 2:
        intro = f"{clients[0]['title']} {clients[0]['name']} & {clients[1]['title']} {clients[1]['name']}"
    else:
        intro = ', '.join([f"{client['title']} {client['name']}" for client in clients[:-1]]) + f", & {clients[-1]['title']} {clients[-1]['name']}"

    # Add the title with proper formatting
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'Efile Confirmation of your {year_str} Tax Declaration\n')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(127, 127, 127)

    # Add the salutation
    salutation = doc.add_paragraph()
    salutation.add_run(f"{intro},")

def add_confirmation_numbers_section(doc, clients, province):
    para = doc.add_paragraph()
    para.add_run("Please be advised that your tax returns were successfully sent via EFile.").bold = True
    para.add_run("\nPlease keep the confirmation numbers for your records:")

    for client in clients:
        isMailQC = False
        # Add the client's title and name
        if len(clients) > 1:
            para.add_run(f"\n\n{client['title']}. {client['name']}:")

        # Iterate over the years to get confirmation numbers
        for year_info in client['years']:
            year = year_info['year']
            confirmation_numbers = year_info['confirmationNumbers']
            cra_confirmation = confirmation_numbers.get("federal", "Not provided")
            if province == "QC":
                qc_confirmation = confirmation_numbers.get("quebec", "Not provided")

            t1135_confirmation = confirmation_numbers.get("t1135", "Not provided")

            # Display confirmation numbers for the specific year
            para.add_run(f"\n{year}:\n").underline = True
            para.add_run(f"CRA: {cra_confirmation}")
            if province == "QC" and qc_confirmation == "Mail QC":
                isMailQC = True
            if province == "QC" and qc_confirmation != "Mail QC":
                para.add_run(f"\nQC: {qc_confirmation}")
            if t1135_confirmation != "Not provided" and t1135_confirmation != "":
                para.add_run(f"\nT1135: {t1135_confirmation}")

        if isMailQC:
          para.add_run("\nPlease do not forget to mail the QC return as explained in the summary.").bold = True

    para = doc.add_paragraph()
    para.add_run('You will receive a notice of assessment from the Canada Revenue ').bold = False
    if province == "QC":
        para.add_run('and Québec Revenue ').bold = False
    para.add_run('to confirm receipt of your tax returns. ').bold = False
    para.add_run('If there is a discrepancy ').bold = True
    para.add_run('between the notice of assessment and the results mentioned in the summary that we emailed you, please send us a copy of the notice of assessment as soon as possible.').bold = False

def add_tax_summary_fees_section(doc, selected_prices, gst_rate, qst_rate, province, includeTaxes):
    subtotal = 0
    percentage_adjustments = 0

    # Calculate subtotal and percentage adjustments
    for price in selected_prices:
        if price['type'] == "number":
            subtotal += price['amount'] * price['quantity']
        elif price['type'] == "%":
            percentage_adjustments += price['amount']

    adjustment_amount = subtotal * (percentage_adjustments / 100)
    adjusted_subtotal = subtotal + adjustment_amount

    if includeTaxes:
        gst = adjusted_subtotal * (gst_rate / 100)
        qst = adjusted_subtotal * (qst_rate / 100) if province == "QC" else 0
    else:
        gst = 0
        qst = 0

    # Calculate total as the adjusted subtotal plus taxes
    total = adjusted_subtotal + gst + qst

    # Add the title for the fees section
    para = doc.add_paragraph("")
    run = para.add_run("Tax Declaration Fees: ")
    run.bold = True
    run.underline = True

    para.add_run(" (Attached is a copy of the invoice for your reference)\n")  # Adding a line break

    # Display the services and their amounts
    for price in selected_prices:
        service_name = price['service']['en']
        if price['type'] == "number":
            amount_display = f"${price['amount']:.2f}"
            if price['quantity'] > 1:
                amount_display += f" x {price['quantity']} = ${(price['amount'] * price['quantity']):.2f}"
        elif price['type'] == "%":
            amount_display = f"{price['amount']:.0f}%"
        else:
            continue

        # Add each service and amount to the same paragraph with line breaks
        run = para.add_run(f"{service_name}: {amount_display}\n")

    # Add subtotal information
    if includeTaxes:
        para.add_run("Subtotal: ")
        para.add_run(f"${adjusted_subtotal:.2f}\n")
        para.add_run(f"GST (837298611RT0001): ${gst:.2f}\n")
        if province == "QC":
            para.add_run(f"QST (1213576115TQ0002): ${qst:.2f}\n")

    total_run = para.add_run(f"Total: ${total:.2f}")
    total_run.bold = True

    # Payment instructions in italic
    payment_para = doc.add_paragraph("")
    payment_run1 = payment_para.add_run("You can pay the fees by Interac e-transfer to ")
    add_hyperlink(payment_para, "taxdeclaration@gmail.com", "mailto:taxdeclaration@gmail.com")
    payment_run2 = payment_para.add_run(". For the password, you can use the word ")
    password_run = payment_para.add_run("declaration")
    password_run.underline = True
    payment_run3 = payment_para.add_run(".")

    # Set all runs to italic
    for run in [payment_run1, payment_run2, password_run, payment_run3]:
        run.italic = True

def add_footer_section(doc, province):
    footer_para = doc.add_paragraph(" ")
    footer_title = footer_para.add_run("Register to your online Tax Account:\n")
    footer_title.bold = True

    footer_para.add_run("It's highly recommended as it allows you to consult tax data, check your notice of assessments, your benefit and credit payments, view your RRSP/TFSA limit, carry forward amounts, and so much more.\n")
    footer_para.add_run("Sign up with")
    add_hyperlink(footer_para, " Canada Revenue Agency", "https://www.canada.ca/en/revenue-agency.html")
    if province == "QC":
        footer_para.add_run("\nSign up with")
        add_hyperlink(footer_para, " Sign up with Quebec Revenue", "https://www.revenuquebec.ca/en/")
    if province == "QC":
        footer_para.add_run("\nTo learn how to register your accounts with CRA and QC Revenue, you can refer to the following videos: ")
    else:
        footer_para.add_run("\nTo learn how to register your accounts with CRA, you can refer to the following video: ")
    footer_para.add_run("\nCanada Revenue: ")
    add_hyperlink(footer_para, "https://youtu.be/_p-lq3eqD6w: ", "https://youtu.be/_p-lq3eqD6w")

    if province == "QC":
        footer_para.add_run("\nQuebec Revenue: ")
        add_hyperlink(footer_para, " https://youtu.be/YuD8i0Le4l4: ", "https://youtu.be/YuD8i0Le4l4")

    if province == "QC":
        footer_para.add_run("\n\nNote that you need to notify Canada Revenue and Quebec Revenue if one of the following cases occurs: \n")
    else:
        footer_para.add_run("\n\nNote that you need to notify Canada Revenue if one of the following cases occurs: \n")
    changes = [
        "   •  Changes in your mailing address",
        "   •  Changes in your bank information",
        "   •  Changes in your marital status",
        "   •  Changes in your province of residency",
        "   •  Becoming a non-Canadian Resident"
    ]

    for change in changes:
        footer_para.add_run(change + "\n")

    footer_para.add_run("\nThank you for using our services & have a profitable year!")

def create_confirmation_email_french(directory_path, clients, selected_prices, tax_rate, includeTaxes):
    province = tax_rate.get("province")
    # Ensure rates are numeric (allow strings like "5" or "5%")
    gst_rate = _parse_rate(tax_rate.get("fedRate", 0))
    qst_rate = _parse_rate(tax_rate.get("provRate", 0)) if province == "QC" else 0

    doc = Document()

    # Set default font to Calibri and size 10pt for the entire document
    set_default_font(doc, "Calibri", 10)

    for client in clients:
        if client['title'] == 'Mr':
            client['title'] = 'M.'
        elif client['title'] in ['Mrs', 'Ms']:
            client['title'] = 'Mme'

    add_intro_section_french(doc, clients)
    add_confirmation_numbers_section_french(doc, clients, province)
    if selected_prices:
        add_tax_summary_fees_section_french(doc, selected_prices, gst_rate, qst_rate, province, includeTaxes)
    add_footer_section_french(doc, province)

    years = sorted(set(year['year'] for client in clients for year in client['years']))
    if len(years) == 2:
        year_str = f"{years[0]} & {years[1]}"
    elif len(years) > 2:
        year_str = ", ".join(str(year) for year in years[:-1]) + f", & {years[-1]}"
    else:
        year_str = str(years[0])
    client_names = " & ".join([f"{client['name']}" for client in clients])
    output_file_path = f"{directory_path}/Confirmation {client_names} {year_str}.docx"

    doc.save(output_file_path)

def add_intro_section_french(doc, clients):
    # Get unique years from all clients
    years = sorted(set(year['year'] for client in clients for year in client['years']))

    # Format the years string in French
    if len(years) == 2:
        year_str = f"{years[0]} & {years[1]}"
    elif len(years) > 2:
        year_str = ", ".join(str(year) for year in years[:-1]) + f", & {years[-1]}"
    else:
        year_str = str(years[0])

    # Format the client introduction based on the number of clients
    if len(clients) == 1:
        intro = f"{clients[0]['title']} {clients[0]['name']}"
    elif len(clients) == 2:
        intro = f"{clients[0]['title']} {clients[0]['name']} & {clients[1]['title']} {clients[1]['name']}"
    else:
        intro = ', '.join([f"{client['title']} {client['name']}" for client in clients[:-1]]) + f", & {clients[-1]['title']} {clients[-1]['name']}"

    # Add the title with proper formatting
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f'Confirmation Efile de votre déclaration d\'impôts {year_str}\n')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(127, 127, 127)

    # Add the salutation
    salutation = doc.add_paragraph()
    salutation.add_run(f"{intro},")

def add_confirmation_numbers_section_french(doc, clients, province):
    para = doc.add_paragraph()
    para.add_run("Veuillez noter que vos déclarations de revenus ont été envoyées avec succès par EFile.").bold = True
    para.add_run("\nVeuillez conserver les numéros de confirmation pour vos dossiers :")

    for client in clients:
        # Add the client's title and name
        if len(clients) > 1:
            para.add_run(f"\n\n{client['title']} {client['name']}:")

        # Iterate over the years to get confirmation numbers
        for year_info in client['years']:
            isMailQC = False
            year = year_info['year']
            confirmation_numbers = year_info['confirmationNumbers']
            cra_confirmation = confirmation_numbers.get("federal", "Non fourni")
            if province == "QC":
                qc_confirmation = confirmation_numbers.get("quebec", "Non fourni")
            t1135_confirmation = confirmation_numbers.get("t1135", "Non fourni")
            # Display confirmation numbers for the specific year
            para.add_run(f"\n{year} :\n").underline = True
            para.add_run(f"ARC: {cra_confirmation}")
            if province == "QC" and qc_confirmation == "Mail QC":
                isMailQC = True
            if province == "QC" and qc_confirmation != "Mail QC":
                para.add_run(f"\nQC: {qc_confirmation}")
            if t1135_confirmation != "Non fourni" and t1135_confirmation != "":
                para.add_run(f"\nT1135: {t1135_confirmation}")

        if isMailQC:
          para.add_run("\nN'oubliez pas de poster la déclaration de QC comme expliqué dans le sommaire.").bold = True

    para = doc.add_paragraph()
    para.add_run("Vous recevrez un avis de cotisation de l'Agence du revenu du Canada ").bold = False
    if province == "QC":
        para.add_run("et de Revenu Québec ").bold = False
    para.add_run("pour confirmer la réception de vos déclarations de revenus. ").bold = False
    para.add_run("S'il y a un écart ").bold = True
    para.add_run("entre l'avis de cotisation et les résultats mentionnés dans le résumé que nous vous avons envoyé par e-mail, veuillez nous envoyer une copie de l'avis de cotisation dès que possible.").bold = False

def add_tax_summary_fees_section_french(doc, selected_prices, gst_rate, qst_rate, province, includeTaxes):
    subtotal = 0
    percentage_adjustments = 0

    # Calculate subtotal and percentage adjustments
    for price in selected_prices:
        if price['type'] == "number":
            subtotal += price['amount'] * price['quantity']
        elif price['type'] == "%":
            percentage_adjustments += price['amount']

    adjustment_amount = subtotal * (percentage_adjustments / 100)
    adjusted_subtotal = subtotal + adjustment_amount

    if includeTaxes:
        raise ValueError (gst_rate)
        gst = adjusted_subtotal * (gst_rate / 100)
        qst = adjusted_subtotal * (qst_rate / 100) if province == "QC" else 0
    else:
        gst = 0
        qst = 0

    # Calculate total as the adjusted subtotal plus taxes
    total = adjusted_subtotal + gst + qst

    # Add the title for the fees section
    para = doc.add_paragraph("")
    run = para.add_run("Frais de Services:")
    run.bold = True
    run.underline = True
    para.add_run(" (Ci-joint la facture pour votre référence)\n")  # Adding a line break

    # Function to format currency for French locale
    def format_currency_french(amount):
        return f"{amount:,.2f}".replace(",", " ").replace(".", ",") + " $"

    # Display the services and their amounts in French format
    for price in selected_prices:
        service_name = price['service']['fr']
        if price['type'] == "number":
            amount_display = format_currency_french(price['amount'])
            if price['quantity'] > 1:
                amount_display += f" x {price['quantity']} = {format_currency_french(price['amount'] * price['quantity'])}"
        elif price['type'] == "%":
            amount_display = f"{price['amount']:.0f}%"
        else:
            continue

        # Add each service and amount to the same paragraph with line breaks
        run = para.add_run(f"{service_name}: {amount_display}\n")

    # Add subtotal information in French format
    if includeTaxes:
        para.add_run("Sous-total: ")
        para.add_run(f"{format_currency_french(adjusted_subtotal)}\n")
        para.add_run(f"TPS (837298611RT0001): {format_currency_french(gst)}\n")
        if province == "QC":
            para.add_run(f"TVQ (1213576115TQ0002): {format_currency_french(qst)}\n")

    total_run = para.add_run(f"Total: {format_currency_french(total)}")
    total_run.bold = True

    # Payment instructions in italic
    payment_para = doc.add_paragraph("")
    payment_run1 = payment_para.add_run("Vous pouvez payer les frais par virement électronique Interac à ")
    add_hyperlink(payment_para, "taxdeclaration@gmail.com", "mailto:taxdeclaration@gmail.com")
    payment_run2 = payment_para.add_run(". Pour le mot de passe, vous pouvez utiliser le mot ")
    password_run = payment_para.add_run("declaration")
    password_run.underline = True
    payment_run3 = payment_para.add_run(".")

    # Set all runs to italic
    for run in [payment_run1, payment_run2, password_run, payment_run3]:
        run.italic = True

def add_footer_section_french(doc, province):
    footer_para = doc.add_paragraph(" ")
    footer_title = footer_para.add_run("Inscrivez-vous à votre compte de taxes en ligne:\n")
    footer_title.bold = True

    footer_para.add_run("C'est fortement recommandé car il vous permet de consulter les données fiscales, de vérifier vos avis de cotisation et vos paiements de prestations et de crédits, de visualiser votre limite REER / CELI, vos montants reportés, etc.\n")
    footer_para.add_run("S'inscrire à")
    add_hyperlink(footer_para, " l'Agence Revenu Canada", "https://www.canada.ca/fr/agence-revenu.html")
    if province == "QC":
        footer_para.add_run("\nS'inscrire à")
        add_hyperlink(footer_para, " Revenu Québec", "https://www.revenuquebec.ca/fr/")
    if province == "QC":
        footer_para.add_run("\nPour savoir comment s'inscrire à Revenu Canada et Revenu Québec, veuillez consulter les vidéos suivantes: ")
    else:
        footer_para.add_run("\nPour savoir comment s'inscrire à Revenu Canada, veuillez consulter la vidéo suivante: ")
    footer_para.add_run("\nRevenu Canada: ")
    add_hyperlink(footer_para, "https://youtu.be/4t6pqHX58H4", "https://youtu.be/4t6pqHX58H4")

    if province == "QC":
        footer_para.add_run("\nRevenu QC: ")
        add_hyperlink(footer_para, "https://youtu.be/e8pY3zSvStM", "https://youtu.be/e8pY3zSvStM")

    footer_para.add_run("\n\nNotez que vous devez aviser Revenu Canada et Revenu Québec si l'un des cas suivants se produit: \n")
    changes = [
        "   •  Changements dans votre adresse postale",
        "   •  Modifications de vos informations bancaires",
        "   •  Changements dans votre état civil",
        "   •  Changements de votre province de résidence",
        "   •  Devenir un non-résident"
    ]

    for change in changes:
        footer_para.add_run(change + "\n")

    footer_para.add_run("\nMerci d'avoir utilisé nos services et ayez une année rentable !")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
