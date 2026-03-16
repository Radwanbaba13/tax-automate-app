"""
Helper utilities for applying doc_text_config to Word document generation.

The config dict structure (from DB / docTextDefaults.ts):
{
  "individual_en": { "docTitle": { "text": "...", "style": { ... } }, ... },
  "couple_en": { ... },
  ...
}

Each block: { text: str, style: { fontSize, color, bold, italic, underline, alignment } }
Text may contain {variable} placeholders substituted at generation time.
"""

import json

from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def is_segmented(text):
    """Return True if text is a segments JSON array (starts with '[')."""
    return isinstance(text, str) and text.strip().startswith('[')


def parse_segments(text):
    """Parse a segments JSON string into a list of segment dicts."""
    try:
        segments = json.loads(text)
        if isinstance(segments, list):
            return segments
    except (json.JSONDecodeError, TypeError):
        pass
    return None


def resolve_doc_type_key(language, is_couple=False, is_multiyear=False):
    """Return the config key like 'individual_en', 'couple_fr', etc."""
    lang = 'en' if language == 'EN' else 'fr'
    if is_multiyear:
        return f'multiyear_{lang}'
    if is_couple:
        return f'couple_{lang}'
    return f'individual_{lang}'


def get_cfg(doc_text_config, doc_type_key):
    """Get the resolved config dict for a specific doc type, or empty dict."""
    if not doc_text_config:
        return {}
    return doc_text_config.get(doc_type_key, {})


def _get_raw_text(cfg, key, default):
    """Get raw text from config block (may be segments JSON). For internal use."""
    block = cfg.get(key, {}) if cfg else {}
    return block.get('text', default) if block else default


def get_text(cfg, key, default, **kwargs):
    """Get text from config block, substitute {variables}, fall back to default.
    If text is segments JSON, flattens it to plain text so callers using
    para.add_run(get_text(...)) never see raw JSON."""
    text = _get_raw_text(cfg, key, default)
    if is_segmented(text):
        segments = parse_segments(text)
        if segments:
            text = ''.join(seg.get('t', '') for seg in segments)
    if kwargs:
        try:
            text = text.format(**kwargs)
        except (KeyError, IndexError):
            pass
    return text


def get_style(cfg, key):
    """Get style dict from config block, or empty dict."""
    block = cfg.get(key, {}) if cfg else {}
    return block.get('style', {}) if block else {}


def _parse_color(color_str):
    """Parse '#rrggbb' into RGBColor."""
    if not color_str:
        return None
    color_str = color_str.lstrip('#')
    if len(color_str) == 6:
        r, g, b = int(color_str[0:2], 16), int(color_str[2:4], 16), int(color_str[4:6], 16)
        return RGBColor(r, g, b)
    return None


def apply_style(run, style):
    """Apply a style dict to a python-docx run. Only sets properties that exist in the dict."""
    if not style:
        return run
    if 'bold' in style:
        run.bold = style['bold']
    if 'italic' in style:
        run.italic = style['italic']
    if 'underline' in style:
        run.underline = style['underline']
    if 'fontSize' in style and style['fontSize']:
        run.font.size = Pt(style['fontSize'])
    if 'color' in style and style['color']:
        rgb = _parse_color(style['color'])
        if rgb:
            run.font.color.rgb = rgb
    return run


def apply_alignment(paragraph, style):
    """Apply alignment from style dict to a paragraph."""
    if not style:
        return
    alignment = style.get('alignment')
    if alignment == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == 'left':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def styled_run(para, cfg, key, default_text, default_style=None, **kwargs):
    """
    Add a run to the paragraph using config text + style (falling back to defaults).
    If the config text is a segments JSON array, creates one run per segment with
    per-segment bold/italic/underline overrides on top of the base style.

    Args:
        para: python-docx paragraph
        cfg: resolved doc type config dict (e.g. config['individual_en'])
        key: block key (e.g. 'docTitle')
        default_text: fallback text if config doesn't have this key
        default_style: dict of default style props (e.g. {'bold': True})
        **kwargs: variables for {placeholder} substitution in text

    Returns:
        The last created run.
    """
    raw_text = _get_raw_text(cfg, key, default_text)
    base_style = get_style(cfg, key) or default_style or {}

    if is_segmented(raw_text):
        segments = parse_segments(raw_text)
        if segments:
            last_run = None
            for seg in segments:
                seg_text = seg.get('t', '')
                if kwargs:
                    try:
                        seg_text = seg_text.format(**kwargs)
                    except (KeyError, IndexError):
                        pass
                run = para.add_run(seg_text)
                seg_style = dict(base_style)
                if 'b' in seg:
                    seg_style['bold'] = seg['b']
                if 'i' in seg:
                    seg_style['italic'] = seg['i']
                if 'u' in seg:
                    seg_style['underline'] = seg['u']
                if 'c' in seg:
                    seg_style['color'] = seg['c']
                apply_style(run, seg_style)
                last_run = run
            return last_run or para.add_run('')

    # Plain text path
    text = raw_text
    if kwargs:
        try:
            text = text.format(**kwargs)
        except (KeyError, IndexError):
            pass
    run = para.add_run(text)
    apply_style(run, base_style)
    return run
